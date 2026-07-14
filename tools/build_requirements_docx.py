from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D9E2EC") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1)
    section.footer_distance = Cm(1)

    styles = doc.styles
    styles["Normal"].font.name = "SimSun"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string("1F2937")
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 24, "0F172A"),
        ("Subtitle", 13, "4B5563"),
        ("Heading 1", 17, "0F172A"),
        ("Heading 2", 13, "1E3A5F"),
    ]:
        style = styles[name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(8)


def add_para(doc: Document, text: str = ""):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74) if text else None
    return p


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(75, 85, 99)


def add_full_figure(doc: Document, img: Path, caption: str, width: Cm = Cm(14.6)) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img), width=width)
    add_caption(doc, caption)


def add_figure_grid(doc: Document, items: list[tuple[Path, str]], cols: int = 2, width: Cm = Cm(7.1)) -> None:
    rows = (len(items) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = False
    idx = 0
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "F8FAFC")
            set_cell_border(cell)
            if idx >= len(items):
                cell.text = ""
                idx += 1
                continue
            img, caption = items[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            p.add_run().add_picture(str(img), width=width)
            cp = cell.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.first_line_indent = None
            for run in cp.runs:
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(75, 85, 99)
            idx += 1
    doc.add_paragraph()


def add_cover(doc: Document, media: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(media / "image1.png"), width=Cm(5.8))
    title = doc.add_paragraph("狭域巡航 aiPotRol")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(6)
    title.runs[0].font.name = "SimSun"
    title.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    title.runs[0].font.size = Pt(22)
    title.runs[0].font.bold = True
    title.runs[0].font.color.rgb = RGBColor.from_string("0F172A")
    subtitle = doc.add_paragraph("仓储环境智能巡检与货物异常监测小车")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.name = "SimSun"
    subtitle.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string("334155")
    doc_type = doc.add_paragraph("需求分析文档")
    doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_type.paragraph_format.space_before = Pt(12)
    doc_type.runs[0].font.name = "SimSun"
    doc_type.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc_type.runs[0].font.size = Pt(18)
    doc_type.runs[0].font.bold = True
    doc_type.runs[0].font.color.rgb = RGBColor.from_string("1E3A5F")
    meta = doc.add_paragraph("版本：V4.1    日期：2026 年 7 月 13 日")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.name = "SimSun"
    meta.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    meta.runs[0].font.size = Pt(10.5)
    meta.runs[0].font.color.rgb = RGBColor.from_string("64748B")
    doc.add_page_break()


def add_alert_record_table(doc: Document) -> None:
    rows = [
        ("时间", "点位", "类别", "等级", "记录内容"),
        ("09:24", "P2 A 区管线左侧", "通道占用", "预警", "检测到纸箱进入通道，建议减速观察并准备手动接管。"),
        ("09:31", "P4 B 区设备段", "视觉异常", "严重", "识别到人员进入巡检区域，触发重点告警并暂停巡检。"),
        ("09:38", "P5 高风险设备区", "环境异常", "严重", "可燃气体数值超过阈值，建议现场复核并排查泄漏风险。"),
        ("09:45", "P6 返回起点", "巡检异常", "提示", "返回起点过程中网络波动，系统保留最新状态并提示重新连接。"),
    ]
    table = doc.add_table(rows=len(rows), cols=5)
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_border(cell)
            if r_idx == 0:
                set_cell_shading(cell, "E8F1FA")
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
            else:
                set_cell_shading(cell, "FFFFFF")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx < 4 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            run = p.add_run(value)
            run.font.name = "SimSun"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string("1F2937")
            if r_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string("1E3A5F")
    p = doc.add_paragraph("表 1 告警记录列表示例：按时间、点位、类别、等级和内容汇总巡检异常。")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(75, 85, 99)


def build(doc: Document, media: Path) -> None:
    add_cover(doc, media)

    doc.add_heading("一、项目背景", level=1)
    add_para(doc, "“狭域巡航 aiPotRol”是在仓储环境智能巡检方向上进一步细化形成的项目。相比传统宽泛的仓储巡检，本项目更关注仓库中容易被忽略、也更适合小车落地的狭窄通道、货架间通道、管廊式仓储区域和高风险设备段。这些区域通常空间有限、视线受遮挡，人工巡查容易遗漏通道堵塞、物体倾倒、人员闯入和环境异常。")
    add_para(doc, "项目以润和 iCar 智能小车为移动平台，结合 HarmonyOS 上位机 APP、TCP 网络通信、视频回传、麦克纳姆轮控制、环境监测、雷达通道安全判断、视觉异常识别和告警记录展示，形成一套能够现场演示、能够继续扩展的巡检系统。当前项目并不是停留在概念说明阶段，代码中已经实现了较完整的 APP 页面、状态模型、通信命令和消息解析逻辑。")
    add_para(doc, "本项目的目标可以概括为三个方面：第一，让小车能够在狭窄仓储通道中被稳定连接、查看和接管；第二，让巡检过程以点位、路线、状态和进度的形式被清楚展示；第三，让环境、雷达、视觉和巡检过程中的异常统一进入告警记录，便于现场处理和后续复盘。")
    add_full_figure(doc, media / "image2.png", "图 1 仓储狭域巡检总览：以管廊/货架通道为核心，展示小车路线、巡检点位和异常区域。", Cm(14.4))

    doc.add_heading("二、项目核心功能", level=1)
    add_para(doc, "项目核心功能围绕“连接小车、控制小车、查看巡检、发现异常、记录告警”展开。APP 首页以“狭域巡航”为入口，显示在线状态、当前点位、下一点位、巡检进度和重点告警，用户可以从首页进入环境感知、SLAM 导航、激光雷达、视觉识别和手动接管等模块。代码中的页面清单已经包括 Index、NetworkSettings、RemoteControl、MecanumWheel、PatrolDashboard、PatrolControl、EnvironmentMonitor、AlertRecords、Module1Environment、Module2Navigation、Module3Lidar、Module4Vision 和 Module5AppMonitor。")
    add_para(doc, "网络连接功能支持配置小车 IP、TCP 指令端口和视频端口。连接成功后，APP 会初始化消息处理器并进入在线同步状态；连接失败时给出提示。")
    add_para(doc, "手动接管功能已经覆盖方向按键、摇杆控制、麦克纳姆轮控制和急停。底盘控制协议中包含自由控制、按键控制、四轮独立速度控制、保存图片、开始录制、结束录制、启动巡航和中止巡航等命令。麦克纳姆轮控制是项目亮点之一，它能体现小车在狭窄通道中的横移、原地调整和精细避让能力。")
    add_para(doc, "自动巡检功能已经在 APP 中形成操作入口，支持手动模式和自动巡检模式切换，并支持开始巡检、暂停巡检、继续巡检、返回起点和停止巡检。巡检路线使用 warehouse_A 和 warehouse_B 两类路线标识，默认点位包括 P1 管廊入口、P2 A 区管线左侧、P3 A 区管线右侧、P4 B 区设备段、P5 高风险设备区和 P6 返回起点。APP 可以显示当前点、下一点、巡检状态和进度，异常发生后也可以将对应点位标记为异常。")
    add_para(doc, "环境监测模块围绕仓储安全常见指标设计，包含温度、湿度、可燃气体、PM2.5、光照、烟雾和气压等数据。代码中已经定义环境快照、环境历史记录和环境阈值，页面可以显示实时数值和最近历史。雷达模块记录前方距离、左侧距离、右侧距离、通道状态和更新时间，通道状态包括通道畅通、减速通行、通道堵塞和通道占用。视觉模块用于表达人员闯入、物体占道、货物异常和红色瓶体演示等内容。")
    add_para(doc, "告警模块已经实现统一记录结构，能够保存告警时间、点位、类别、消息、等级和图片地址，并预留 AI 分析结果。告警类别包括环境异常、通道占用、视觉异常和巡检异常。告警记录页支持查看单条告警分析，也支持围绕告警数据进行问答式说明，这部分让项目在异常处置上不只是提示“有问题”，还能进一步给出风险和处理建议。")
    add_figure_grid(doc, [
        (media / "image3.png", "图 2 网络连接设置：配置小车 IP、指令端口和视频端口，支持在线连接与离线预览。"),
        (media / "image4.png", "图 3 APP 首页总览：展示在线状态、巡检点位、重点告警和各功能模块入口。"),
    ], cols=2, width=Cm(5.2))
    add_figure_grid(doc, [
        (media / "image8.png", "图 4 手动接管页面：方向按键、急停和视频画面用于现场快速干预。"),
        (media / "image9.png", "图 5 摇杆控制页面：通过虚拟摇杆实现更连续的小车方向控制。"),
    ], cols=2, width=Cm(7.1))
    add_full_figure(doc, media / "image10.png", "图 6 麦克纳姆轮控制：支持四轮速度调节、横向移动、原地旋转和精细避让。", Cm(14.4))

    doc.add_heading("三、应用场景说明", level=1)
    add_para(doc, "系统的典型路线不是完整大仓库盘点，而是从入口进入管廊或货架通道，依次经过若干固定点位，在重点区域停留观察，并在发现异常后把异常记录到 APP。")
    add_para(doc, "第一类场景是狭窄通道巡检。小车从入口出发，沿 P1 到 P6 点位移动，APP 展示当前点、下一点和巡检进度。第二类场景是通道占用。用户可以在通道中放置纸箱、瓶体或其他物体，雷达模块显示通道占用或堵塞，视觉模块记录物体异常。第三类场景是人员闯入。人员进入警戒区或巡检通道后，视觉模块产生人员异常记录，APP 告警页展示时间、点位和等级。第四类场景是环境异常。通过真实传感器或模拟数据触发温度、可燃气体、PM2.5、光照等异常，环境模块和告警模块同步展示。第五类场景是人工接管。自动巡检遇到异常时，用户可以切换手动模式，结合视频画面控制小车停止、后退、横移或返回起点。")
    add_figure_grid(doc, [
        (media / "image5.png", "图 7 点位化巡检路线：以 P1-P6 固定点组织任务，便于定位、演示和复盘。"),
        (media / "image6.png", "图 8 RViz 建图效果：展示地图轮廓，可用于后续路线规划与导航调试。"),
    ], cols=2, width=Cm(6.8))
    add_full_figure(doc, media / "image7.png", "图 9 巡检点位与异常区域标注：在地图上直观呈现路线、点位和风险区域。", Cm(12.8))

    doc.add_heading("四、功能性说明", level=1)
    doc.add_heading("4.1 上位机 APP 功能", level=2)
    add_para(doc, "APP 是项目最主要的展示和操作入口。首页需要承担总览作用，用户打开应用后应能直接看到项目名称、设备在线状态、当前巡检点位、下一巡检点位、巡检进度和近期重点告警。当前 Index 页面已经实现这些内容，并且通过模块卡片把环境感知、SLAM 导航、激光雷达、视觉识别和手动接管串联起来。")
    add_para(doc, "网络设置功能用于真实小车联调。用户可以配置默认 IP、指令端口和视频端口，连接成功后进入在线同步状态，连接失败时给出提示。APP 保留离线预览模式，用于在没有真实小车或网络环境不稳定时展示页面结构和业务流程。")
    add_para(doc, "巡检监控功能用于展示小车自动巡检的实时状态。页面显示路线名称、当前点位、下一点位、巡检状态、巡检进度、雷达状态、视觉状态和最新告警。PatrolDashboard 页面已经按这个思路实现，并使用点位卡片展示路线中的已巡检点和异常点。")
    add_para(doc, "巡检控制功能用于下发自动巡检命令。用户可以在手动遥控和自动巡检之间切换，可以启动巡检、暂停巡检、继续巡检、停止巡检和返回起点。当前 PatrolControl、Module2Navigation 和 Module5AppMonitor 页面都接入了巡检控制工具，底层通过 JSON 命令向小车端发送 start_patrol、pause_patrol、resume_patrol、return_home 和 stop_patrol。")
    doc.add_heading("4.2 小车运动与手动接管功能", level=2)
    add_para(doc, "手动接管是项目安全性的关键。小车在自动巡检过程中如果遇到通道堵塞、人员靠近、识别异常或导航不稳定，用户应能够立即切换到手动模式。当前项目已经实现方向按键、摇杆控制、麦克纳姆轮控制和急停入口，能够覆盖前进、后退、左平移、右平移、左旋转、右旋转、刹车停止等基本动作。")
    add_para(doc, "麦克纳姆轮控制是狭域巡航场景的重要支撑。普通两轮或四轮小车在货架通道内调头和避让不方便，而麦克纳姆轮可以横向移动，也可以原地调整姿态，更适合狭窄通道、管廊式区域和货架间通道。项目中的 MecanumWheel 页面和四轮独立速度控制协议为这一能力提供了展示入口。")
    doc.add_heading("4.3 巡检路线、点位与任务状态功能", level=2)
    add_para(doc, "巡检路线采用点位化方式组织，不追求复杂的仓库全域调度，而是围绕可演示、可定位、可复盘的 P 点路线展开。每个点位都包含是否已巡检和是否异常两个状态。任务状态包括空闲、巡检中、已暂停、正在导航、避障中、异常告警中和返回起点。APP 接收到小车端 patrol 类型消息后，会更新当前点、下一点、进度、路线和状态说明。")
    add_para(doc, "雷达检测到通道堵塞或占用时，车辆状态可切换为避障中；告警消息指示 alerting 时，车辆状态可切换为异常告警中。这样的状态设计让巡检过程不只是按钮操作，而是能够表现出小车正在做什么、下一步去哪里、当前是否安全。")
    doc.add_heading("4.4 环境监测功能", level=2)
    add_para(doc, "环境监测面向仓储安全中的常见隐患。系统展示温度、湿度、可燃气体、PM2.5、光照、烟雾和气压等指标，其中温度、湿度、可燃气体、PM2.5 和光照已经在模型中定义阈值。环境数据不仅显示当前值，还形成历史记录，便于说明异常是在巡检过程中发现的，而不是孤立的一次展示。")
    add_para(doc, "环境异常采用正常、预警、严重异常三个等级。轻微超限时 APP 可以作为预警显示，严重超限时进入重点告警。环境记录应绑定点位，例如“P5 高风险设备区可燃气体异常”或“P3 A 区管线右侧光照不足”，这样在后续复盘时能够直接定位到异常区域。")
    add_figure_grid(doc, [
        (media / "image11.png", "图 10 温度监测页面：展示温度曲线、实时数值和阈值状态。"),
        (media / "image12.png", "图 11 湿度监测页面：展示湿度变化趋势和当前安全状态。"),
        (media / "image13.png", "图 12 可燃气体监测页面：用于识别高风险设备区的气体异常。"),
        (media / "image14.png", "图 13 PM2.5 监测页面：展示空气质量指标和巡检点位关联。"),
        (media / "image15.png", "图 14 光照监测页面：用于判断通道或设备区照明不足风险。"),
        (media / "image16.png", "图 15 环境快照记录：汇总多点位环境指标，便于巡检复盘。"),
    ], cols=2, width=Cm(5.1))
    doc.add_heading("4.5 雷达通道安全功能", level=2)
    add_para(doc, "雷达通道安全功能主要用于判断小车前方和左右两侧是否具备安全通行条件。系统记录前方距离、左侧距离和右侧距离，并根据距离判断通道畅通、需要减速、通道堵塞或通道被占用。对于狭窄仓储通道，这一功能可以解释小车为什么需要低速巡检，也可以配合手动接管说明安全策略。")
    add_para(doc, "当前代码已经定义 LidarSnapshot 和 ChannelStatus，并在消息处理器中支持 lidar 类型消息。收到通道堵塞或通道占用状态后，系统会同步更新雷达状态，并把车辆状态切换为避障中。后续接入真实雷达数据后，可以直接沿用现有页面和状态结构。")
    doc.add_heading("4.6 视觉识别与演示功能", level=2)
    add_para(doc, "视觉识别功能用于展示人员闯入、物体占道、货物异常和特定目标检测。当前模型结构已经包含人员检测、箱体或物体进入通道、图像地址和检测说明。Module4Vision 页面还接入了红色瓶体演示相关命令，包括启动演示、停止演示、急停演示、停止蜂鸣器和获取演示状态。")
    add_para(doc, "红色瓶体演示适合作为答辩中的视觉识别样例。系统可以通过 6600 端口请求演示状态，读取距离阈值、速度、是否检测到目标、蜂鸣器是否激活和更新时间。这个功能说明项目并不只是做了静态页面，而是已经为视觉识别和小车动作联动预留了实际控制链路。")
    doc.add_heading("4.7 告警记录与 AI 辅助分析功能", level=2)
    add_para(doc, "告警记录用于把环境、雷达、视觉和巡检过程中的异常统一保存。每条告警包含编号、时间、点位、类别、消息、等级、图片地址和可选 AI 分析结果。首页只展示重点告警，详细记录在 AlertRecords 页面中查看。告警发生后，系统还会尝试触发 AI 分析，生成摘要、风险和处理建议。")
    add_para(doc, "AI 辅助分析不是项目运行的硬性依赖，但它能增强展示效果。即使分析失败，系统也会给出网络异常提示，不影响告警记录本身。这样的设计保证了基础巡检功能稳定，同时为后续智能化处理留下扩展空间。")
    add_alert_record_table(doc)
    add_figure_grid(doc, [
        (media / "image18.png", "图 16 单条告警分析：展示异常详情、风险说明和处理建议。"),
        (media / "image19.png", "图 17 AI 分析结果页：对告警原因、影响范围和处置措施进行结构化说明。"),
        (media / "image20.png", "图 18 AI 分析触发状态：展示分析过程、提示信息和失败兜底。"),
        (media / "image21.png", "图 19 告警复盘页面：集中展示多条异常及其处理状态。"),
        (media / "image22.png", "图 20 告警问答页面：围绕巡检异常进行追问和解释。"),
        (media / "image23.png", "图 21 AI 告警问答示例：支持按点位、风险等级和处置建议生成回答。"),
    ], cols=2, width=Cm(5.1))

    doc.add_heading("五、非功能性需求", level=1)
    non_func = [
        "在易用性方面，系统应尽量让仓库管理员通过 APP 完成主要操作。连接小车、查看状态、启动巡检、暂停巡检、手动接管、查看告警都应在移动端完成，不要求普通用户进入 Linux 终端或 ROS 调试界面。页面文字应接近现场业务表达，例如当前点位、下一点位、通道占用、人员闯入、可燃气体异常等，避免只显示底层技术词。",
        "在实时性方面，手动控制指令需要尽量短延迟，尤其是急停、刹车和停止巡检。视频画面应能够辅助用户判断小车前方环境，环境数据、雷达状态和视觉状态应在收到消息后及时刷新。告警记录应优先展示严重异常，避免重要信息被普通日志淹没。",
        "在可靠性方面，系统要允许真实连接失败时继续进行页面展示。APP 已经支持离线预览，这是课堂演示和项目答辩中的必要保障。通信层解析 JSON 消息时应容忍分包和残留数据，当前 PatrolMessageHandler 已通过缓冲区方式处理按换行分割的消息。对于未知消息，系统不会直接中断运行。",
        "在安全性方面，手动接管、急停和停止巡检的优先级应高于普通巡检动作。小车处于狭窄通道时应优先保证低速和可控，通道堵塞、人员闯入和严重环境异常都应能够触发告警。对于红色瓶体演示和蜂鸣器控制，系统提供停止演示、急停演示和停止蜂鸣器入口，避免演示状态失控。",
        "在可维护性方面，项目代码已经按页面、组件、通信、状态、模型和工具类分层。页面位于 pages 目录，视频和通用 UI 位于 components 目录，TCP 连接和巡检命令位于 tcp 目录，状态集中在 PatrolStore，数据结构集中在 PatrolModels。这样的结构便于后续继续接入真实传感器、替换视觉模型、调整页面和扩展任务报告。",
        "在可扩展性方面，当前系统已经保留了 SLAM 导航、环境阈值、雷达通道状态、视觉检测日志、AI 告警分析、图片地址和任务状态等接口。后续可以在不大改 APP 页面结构的前提下，接入真实 ROS 节点、YOLO 推理结果、任务报告导出和历史数据查询。",
        "在数据一致性与可追溯方面，巡检任务、点位状态、环境快照、雷达状态、视觉日志和告警记录应围绕同一套点位编号组织。告警信息需要保留发生时间、点位、类别、等级和处置说明，便于从首页快速定位到详情页，也便于后续生成巡检报告。对于同一异常源，系统应避免重复刷屏式告警，优先展示最新状态和最严重等级。",
        "在演示与部署适应性方面，项目需要兼顾真实小车联调和课堂答辩展示。真实部署时应支持配置 IP 与端口，适应不同局域网环境；演示时应允许使用模拟数据、离线页面和预置告警，保证没有小车或传感器时也能完整说明业务闭环。页面在手机端展示时应保持按钮可点、文字可读、截图区域清晰。",
    ]
    for text in non_func:
        add_para(doc, text)

    doc.add_heading("六、概要设计", level=1)
    overview = [
        "系统整体采用上位机 APP、通信控制层、状态管理层、感知与算法层、小车硬件层的分层设计。上位机 APP 面向用户，负责页面展示和操作入口；通信控制层负责 TCP 连接、命令发送和消息解析；状态管理层负责保存车辆、环境、雷达、视觉、告警和点位数据；感知与算法层负责雷达、视觉、环境阈值和后续 SLAM 导航；硬件层负责底盘运动、传感器采集、视频回传和蜂鸣器等外设控制。",
        "APP 端页面结构清晰：首页负责总览，NetworkSettings 负责连接配置，RemoteControl 和 MecanumWheel 负责手动控制，PatrolDashboard 和 PatrolControl 负责巡检监控与任务控制，EnvironmentMonitor 和 Module1Environment 负责环境数据，Module3Lidar 负责通道安全，Module4Vision 负责视觉识别和红色瓶体演示，AlertRecords 负责告警记录和 AI 分析。页面之间围绕“首页总览-模块详情-异常复盘”的路径组织，便于用户从状态发现问题，再进入具体模块处理问题。",
        "通信协议分为两类。一类是底盘控制协议，使用帧格式发送自由控制、按键控制、四轮速度、保存图片、视频录制和巡航控制等命令；另一类是巡检业务 JSON 协议，用于发送 start_patrol、pause_patrol、resume_patrol、return_home、stop_patrol、set_mode 等任务指令，并接收 env、patrol、alert、lidar、vision 和 status 等类型的状态消息。两类协议分工明确，底盘协议保证动作控制，业务协议保证页面状态和告警信息同步。",
        "状态管理采用集中式 Store。PatrolStore 保存当前环境快照、环境历史、小车状态、雷达状态、视觉状态、视觉日志、巡检点位和告警记录。页面通过订阅 Store 刷新视图，通信层收到消息后更新 Store。这样可以减少页面之间的数据重复，也便于在未来加入本地持久化、历史查询或任务报告导出功能。",
        "巡检过程可以概括为：用户连接小车，选择手动或自动模式；自动模式下启动巡检路线，小车按点位推进；环境、雷达、视觉模块持续或按点位返回状态；出现异常时写入告警记录并更新点位状态；用户根据视频和告警决定继续巡检、暂停巡检、返回起点或手动接管。该流程把“巡检-发现-记录-分析-处置”串成闭环。",
        "环境异常、雷达通道占用、视觉识别异常和巡检任务异常统一进入告警记录。告警生成后，首页展示重点告警，详情页展示完整信息，AI 分析模块给出原因、风险和处理建议。急停、停止巡检和手动接管属于高优先级操作，应在异常出现时始终保持可用。",
        "后续可以把当前演示数据逐步替换为真实传感器和算法输出：环境数据由传感器采集，雷达状态由距离或点云算法判断，视觉日志由目标检测模型生成，巡检路线由 SLAM 建图和导航节点提供。由于 APP 已经把页面、协议、状态和模型拆开，新增能力主要集中在通信适配和数据模型补充，不需要重做整体交互结构。",
    ]
    for text in overview:
        add_para(doc, text)


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: build_requirements_docx.py <media_dir> <output_docx>")
        return 2
    media = Path(sys.argv[1])
    output = Path(sys.argv[2])
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    build(doc, media)
    doc.save(output)
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
