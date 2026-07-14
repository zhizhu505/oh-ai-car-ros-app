from __future__ import annotations

import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document


def para_has_image(paragraph) -> bool:
    for run in paragraph.runs:
        if run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict"):
            return True
    return False


def main() -> int:
    path = Path(sys.argv[1])
    doc = Document(path)
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    styles = Counter(p.style.name if p.style is not None else "" for p in doc.paragraphs)
    print("styles_top=", styles.most_common(20))
    print("\nHEADINGS / KEY PARAGRAPHS")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().replace("\n", " ")
        style = p.style.name if p.style is not None else ""
        if text and (
            style.startswith("Heading")
            or "标题" in style
            or "非功能" in text
            or "概要设计" in text
            or "项目说明书" in text
            or "图" in text[:3]
        ):
            print(f"{i:04d} [{style}] {text[:160]}")
    print("\nIMAGE PARAGRAPHS")
    for i, p in enumerate(doc.paragraphs):
        if para_has_image(p):
            before = doc.paragraphs[i - 1].text.strip() if i > 0 else ""
            after = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
            print(f"{i:04d} before={before[:70]!r} after={after[:70]!r}")
    with zipfile.ZipFile(path) as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
    print(f"\nmedia_files={len(media)}")
    for name in media[:50]:
        print(name)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
