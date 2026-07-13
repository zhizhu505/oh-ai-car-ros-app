from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "仓储环境智能巡检与货物异常监测小车-项目需求说明书V3.1.docx"
LOGO = Path(r"E:\xwechat_files\wxid_n6e04glg4yi422_80a4\temp\RWTemp\2026-07\b7928310b85d18969dd41e30e9fbd211.png")
IMG = ROOT / "generated_assets" / "ref_images"
PROTO = ROOT / "doc" / "prototype"


def east_asia(run, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    east_asia(r)
    r.font.size = Pt(10)
    r.bold = bold


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, True, WD_ALIGN_PARAGRAPH.CENTER)
        shade(t.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_text(cells[i], value)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    east_asia(r)
    r.font.size = Pt(10.5)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        east_asia(r)
        r.font.size = Pt(10.5)


def figure(doc, path, caption, width):
    path = Path(path)
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    east_asia(r)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(89, 89, 89)


def heading(doc, text, level):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    east_asia(r)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        r.font.size = Pt(16)
    elif level == 2:
        r.font.size = Pt(13)
    else:
        r.font.size = Pt(11)


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)


def cover(doc):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(2.2))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("仓储环境智能巡检与货物异常监测小车")
    east_asia(r)
    r.bold = True
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目需求说明书")
    east_asia(r)
    r.bold = True
    r.font.size = Pt(18)

    doc.add_paragraph()
    table(
        doc,
        ["项目", "内容"],
        [
            ("文档版本", "V3.1"),
            ("编制日期", "2026 年 7 月 13 日"),
            ("项目定位", "面向仓储狭窄通道、货架通道和管廊式区域的巡检与异常监测小车"),
            ("当前实现", "HarmonyOS 控制 APP、小车网络连接、视频回传、手动接管、巡检控制、环境与告警展示"),
        ],
        [3.5, 11],
    )
    doc.add_page_break()


def build():
    doc = Document()
    setup(doc)
    cover(doc)

    heading(doc, "一、项目背景", 1)
    para(
        doc,
        "本项目最初面向仓储环境巡检。结合实际完成情况和小车硬件特点后，项目场景进一步细化为仓储狭窄通道、货架间通道、管廊式仓储区域的移动巡检。"
        "这类区域通常空间窄、视线受遮挡，人工巡查容易漏掉通道堵塞、物品倾倒、人员闯入和环境异常等问题。"
    )
    para(
        doc,
        "项目采用润和 iCar 智能小车作为移动平台，配合雷达、摄像头、ROS 建图导航能力以及 HarmonyOS 上位机 APP，"
        "实现从小车连接、远程控制、视频查看、巡检任务控制到异常告警展示的一套完整演示系统。"
    )
    para(
        doc,
        "当前版本重点解决“看得见、控得住、讲得清”三个问题：管理端能看到小车状态和视频画面，能在自动巡检和手动接管之间切换，"
        "也能把环境、通道、视觉等异常和具体巡检点位联系起来。"
    )
    figure(doc, IMG / "ref_1.png", "图 1 仓储狭域/管廊式巡检场景示意", 6.4)

    heading(doc, "二、项目核心功能", 1)
    para(
        doc,
        "项目核心功能围绕小车巡检过程展开。管理员通过 APP 连接小车，选择巡检路线或手动控制小车移动；小车沿预设点位巡检，"
        "在关键位置采集环境、通道和视觉信息；发现异常后，APP 显示告警内容、点位、时间和处理提示。"
    )
    table(
        doc,
        ["核心功能", "说明", "当前完成情况"],
        [
            ("小车连接与状态同步", "APP 支持填写 IP、TCP 端口和视频端口，连接后同步在线状态。", "已完成，支持离线预览。"),
            ("视频回传", "在控制页展示小车实时视角，方便判断前方障碍和巡检环境。", "已集成视频组件。"),
            ("手动接管", "支持摇杆、方向按键、麦克纳姆轮控制和急停。", "已完成主要控制入口。"),
            ("自动巡检控制", "支持开始、暂停、继续、停止、返回起点等巡检操作。", "已完成 APP 指令封装。"),
            ("巡检点位展示", "显示当前点位、下一点位、巡检进度，异常点可突出显示。", "已完成 P0-P6 点位展示。"),
            ("环境监测", "展示温湿度、烟雾/可燃气体、光照、PM2.5 等环境数据。", "已完成页面与模拟/接入接口。"),
            ("通道安全监测", "识别通道占用、障碍过近等情况，辅助小车安全通行。", "已完成展示逻辑，后续加强真实雷达联动。"),
            ("视觉异常记录", "识别人员闯入、物体占道、货物倾倒等异常，并形成日志。", "已完成页面和记录结构，后续加强模型识别。"),
            ("告警记录与分析", "按等级、类别、点位保存告警，可补充 AI 分析建议。", "已完成告警列表和分析字段。"),
        ],
        [3.3, 7, 5],
    )
    figure(doc, PROTO / "Index.png", "图 2 APP 首页：状态、模块入口和重点告警", 2.8)
    figure(doc, PROTO / "RemoteControl1.png", "图 3 手动接管页面：视频与遥控控制", 2.8)

    heading(doc, "三、应用场景说明", 1)
    para(
        doc,
        "本项目不把仓库所有业务都放进系统，而是选择最适合小车实物演示的几类场景。这样做的好处是场地容易搭建，功能容易验证，"
        "也能突出小车在狭窄空间内灵活移动和实时发现异常的特点。"
    )
    table(
        doc,
        ["场景", "触发方式", "系统表现"],
        [
            ("狭窄通道巡检", "APP 启动自动巡检路线", "小车按 P0-P6 点位移动，APP 显示当前点、下一点和进度。"),
            ("通道占用", "在通道中放置纸箱、瓶体或其他物体", "系统记录通道异常，告警绑定到当前巡检点。"),
            ("人员闯入", "人员进入警戒区或巡检通道", "视觉模块记录人员异常，APP 展示告警。"),
            ("环境异常", "模拟温度、烟雾、可燃气体或光照异常", "环境卡片进入预警或严重状态，并生成记录。"),
            ("人工接管", "自动巡检遇到障碍或演示需要", "切换手动模式，通过摇杆或按键控制小车。"),
        ],
        [3.5, 5.5, 7],
    )
    figure(doc, IMG / "ref_3.png", "图 4 点位化巡检路线示意", 4.4)

    heading(doc, "四、功能性需求", 1)
    para(doc, "功能性需求按实际使用过程描述，不单独设置复杂编号。每一项需求都应能对应到 APP 页面、通信接口或小车演示动作。")

    heading(doc, "4.1 上位机 APP 需求", 2)
    bullets(
        doc,
        [
            "APP 启动后应提供小车连接入口，支持配置 IP 地址、指令端口和视频端口。",
            "APP 首页应展示小车在线状态、当前巡检点、下一巡检点、巡检进度和最新告警。",
            "APP 应提供环境感知、SLAM 导航、激光雷达、视觉识别和手动接管等模块入口。",
            "APP 应支持自动巡检控制，包括开始、暂停、继续、停止和返回起点。",
            "APP 应支持手动遥控，包括方向按键、摇杆控制、麦克纳姆轮控制和急停。",
            "APP 应提供告警记录页面，能查看异常类型、异常点位、时间、等级和处理说明。",
        ],
    )

    heading(doc, "4.2 小车控制需求", 2)
    bullets(
        doc,
        [
            "小车应支持前进、后退、左平移、右平移、左旋转、右旋转和刹车停止。",
            "麦克纳姆轮控制应能体现小车在狭窄通道中的横移和原地调整能力。",
            "手动控制指令应优先于普通自动巡检动作，急停指令优先级最高。",
            "自动巡检过程中应保留人工接管入口，避免异常情况下小车继续运动。",
        ],
    )

    heading(doc, "4.3 巡检与点位需求", 2)
    bullets(
        doc,
        [
            "系统应支持一条主巡检路线，并可扩展支线巡检路线。",
            "巡检路线应由 P0、P1、P2 等清晰点位组成，点位名称应便于讲解和记录。",
            "巡检过程中应显示当前点、下一点和任务进度。",
            "某点位发生异常后，应能在路线或记录中标记为异常点。",
        ],
    )

    heading(doc, "4.4 环境、雷达与视觉需求", 2)
    bullets(
        doc,
        [
            "环境模块应展示温度、湿度、烟雾/可燃气体、光照、PM2.5 等指标，并支持异常等级判断。",
            "雷达模块应用于通道安全判断，重点关注前方障碍、左右安全距离和通道占用。",
            "视觉模块应识别或记录人员闯入、物体占道、货物倾倒等异常情况。",
            "环境、雷达、视觉产生的异常都应统一进入告警记录，便于后续查看。",
        ],
    )

    heading(doc, "五、非功能性需求", 1)
    table(
        doc,
        ["类别", "要求"],
        [
            ("易用性", "主要操作应在 APP 内完成，不要求管理员使用 Linux 终端。页面文字应直接、清楚，适合现场演示。"),
            ("实时性", "手动控制响应应尽量短，视频画面应能辅助判断现场情况。环境和告警状态应及时刷新。"),
            ("可靠性", "网络连接失败、视频不可用或小车离线时，APP 应给出明确提示，不应误导用户。"),
            ("安全性", "急停、停止巡检和手动接管应优先处理。小车在狭窄通道中应低速运行，避免碰撞。"),
            ("可维护性", "APP 页面、通信工具、状态存储和数据模型应分开维护，方便后续接入真实传感器和算法。"),
            ("可扩展性", "系统应保留 SLAM 导航、YOLO 识别、任务报告导出、历史数据查询等扩展空间。"),
        ],
        [3, 12],
    )

    heading(doc, "六、概要设计", 1)
    para(
        doc,
        "系统整体采用“APP 上位机 + 小车通信服务 + ROS/算法模块 + 底盘与传感器”的结构。APP 负责人机交互和状态展示；"
        "通信层负责把 APP 指令发送到小车，并把小车状态、环境数据、视觉结果和告警消息回传；小车端负责运动控制、传感器采集、建图导航和异常识别。"
    )
    table(
        doc,
        ["层次", "组成", "主要职责"],
        [
            ("交互层", "HarmonyOS APP", "连接配置、视频显示、巡检控制、环境数据、告警记录。"),
            ("通信层", "TCP 指令通信、视频流接口", "指令下发、状态回传、视频画面接入。"),
            ("业务层", "巡检状态、点位管理、告警管理", "维护巡检路线、进度、异常记录和页面状态。"),
            ("算法层", "SLAM、雷达检测、视觉识别、环境阈值判断", "实现定位、通道安全判断和异常识别。"),
            ("硬件层", "iCar 小车、麦克纳姆轮、雷达、摄像头、Jetson/MCU", "完成移动、采集、计算和底盘执行。"),
        ],
        [3, 5.5, 7],
    )
    figure(doc, IMG / "ref_4.png", "图 5 RViz 建图效果，可用于路线规划和点位标注", 6.2)
    figure(doc, IMG / "ref_5.png", "图 6 巡检点位与异常区域标注", 5.4)

    heading(doc, "七、接口与数据说明", 1)
    para(
        doc,
        "当前项目通信以局域网连接为主。APP 通过 TCP 端口向小车发送控制指令，通过视频端口接收实时画面。"
        "底盘控制指令沿用项目中的帧格式，巡检状态、环境数据和告警信息可使用结构化消息扩展。"
    )
    table(
        doc,
        ["接口", "说明"],
        [
            ("TCP 6000", "小车控制和巡检任务指令端口。"),
            ("视频端口 6500", "小车实时画面回传端口。"),
            ("按键控制", "用于前进、后退、平移、旋转、刹车。"),
            ("摇杆控制", "用于连续方向控制。"),
            ("麦轮控制", "用于四轮独立速度控制和狭窄空间调姿。"),
            ("巡检控制", "用于启动巡航、中止巡航、暂停、继续和返回起点。"),
        ],
        [4, 11],
    )

    heading(doc, "八、验收与展示重点", 1)
    table(
        doc,
        ["展示项", "通过标准"],
        [
            ("logo 与项目说明", "文档和 APP 能体现项目主题，讲清楚为什么选择仓储狭域巡检。"),
            ("APP 首页", "能看到在线状态、当前点位、下一点位、进度和告警入口。"),
            ("连接配置", "能配置 IP、指令端口和视频端口，支持真实连接或离线预览。"),
            ("手动遥控", "能控制小车完成移动、横移、旋转和停止。"),
            ("自动巡检", "能下发开始、暂停、继续、停止等巡检命令。"),
            ("异常演示", "能展示通道占用、人员闯入、环境异常等至少三类记录。"),
            ("安全接管", "自动巡检中可以切换手动模式或急停。"),
        ],
        [4, 11],
    )

    heading(doc, "九、后续完善方向", 1)
    bullets(
        doc,
        [
            "进一步接入真实环境传感器，减少模拟数据比例。",
            "补充仓储异常图片数据集，训练更贴合本项目的视觉识别模型。",
            "完善 SLAM 地图保存、点位配置和自动巡检路线管理。",
            "增加巡检报告导出功能，把环境数据、异常截图和巡检结果汇总成单次任务报告。",
            "优化视频延迟和断线重连，提升现场演示稳定性。",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("仓储环境智能巡检与货物异常监测小车 项目需求说明书 V3.1")
    east_asia(r)
    r.font.size = Pt(9)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
