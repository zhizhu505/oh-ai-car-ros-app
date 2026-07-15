from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "狭域巡航aiPotRol-项目说明书V4.1.docx"
LOGO = ROOT / "entry" / "src" / "main" / "resources" / "base" / "media" / "patrol_logo.jpg"
IMG = ROOT / "generated_assets" / "ref_images"
PROTO = ROOT / "doc" / "prototype"
PLACEHOLDER_DIR = ROOT / "generated_assets" / "project_doc_placeholders_v41"


PLACEHOLDERS = [
    ("01_network_settings.png", "网络连接页截图占位", "对应页面：pages/NetworkSettings；用于展示 IP、TCP 端口、视频端口和连接状态。"),
    ("02_patrol_dashboard.png", "巡检监控页截图占位", "对应页面：pages/PatrolDashboard；用于展示路线、点位、进度、雷达状态、视觉状态和最新告警。"),
    ("03_patrol_control.png", "巡检控制页截图占位", "对应页面：pages/PatrolControl；用于展示手动/自动模式、开始、暂停、继续、返航和停止。"),
    ("04_environment_monitor.png", "环境监测页截图占位", "对应页面：pages/EnvironmentMonitor、pages/Module1Environment；用于展示温湿度、可燃气体、PM2.5、光照等数据。"),
    ("05_lidar_channel.png", "雷达通道安全页截图占位", "对应页面：pages/Module3Lidar；用于展示前方距离、左右距离、通道占用和避障状态。"),
    ("06_vision_detection.png", "视觉识别页截图占位", "对应页面：pages/Module4Vision；用于展示人员闯入、物体占道、红色瓶体演示和蜂鸣器控制。"),
    ("07_alert_records.png", "告警记录页截图占位", "对应页面：pages/AlertRecords；用于展示环境、雷达、视觉、巡检异常以及 AI 分析结果。"),
    ("08_demo_scene.png", "实物演示场景照片占位", "建议替换为最终小车、仓储狭窄通道、障碍物和人员闯入演示照片。"),
]


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    set_font(r, 16 if level == 1 else 13 if level == 2 else 11, True, "000000")
    return p


def para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_font(r, 10.5)
    return p


def numbered(doc, items):
    for idx, text in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(f"{idx}. {text}")
        set_font(r, 10.5)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, 9, False, "666666")


def figure(doc, path, text, width=5.2):
    path = Path(path)
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, text)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_font(r, 10, True)
        shade(c, "D9EAF7")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_font(r, 10)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return t


def make_placeholders():
    PLACEHOLDER_DIR.mkdir(parents=True, exist_ok=True)
    try:
        title_font = ImageFont.truetype("C:/Windows/Fonts/simhei.ttf", 34)
        body_font = ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", 22)
        small_font = ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", 18)
    except Exception:
        title_font = body_font = small_font = None
    for filename, title, note in PLACEHOLDERS:
        path = PLACEHOLDER_DIR / filename
        img = Image.new("RGB", (1200, 760), "white")
        draw = ImageDraw.Draw(img)
        draw.rectangle((30, 30, 1170, 730), outline=(48, 84, 120), width=4)
        draw.rectangle((30, 30, 1170, 120), fill=(217, 234, 247), outline=(48, 84, 120), width=4)
        draw.text((60, 55), title, fill=(0, 0, 0), font=title_font)
        draw.text((60, 170), f"图片文件名：{filename}", fill=(0, 0, 0), font=body_font)
        draw.text((60, 225), note, fill=(60, 60, 60), font=body_font)
        draw.text((60, 640), "替换方式：保持文件名不变，用最终截图或实物照片覆盖此占位图。", fill=(130, 40, 40), font=small_font)
        img.save(path)


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.add_run().add_picture(str(LOGO), width=Inches(3.0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("狭域巡航 aiPotRol")
    set_font(r, 24, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("仓储环境智能巡检与货物异常监测小车")
    set_font(r, 18, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目说明书")
    set_font(r, 20, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("版本：V4.1    日期：2026 年 7 月 13 日")
    set_font(r, 11)
    doc.add_page_break()


def build():
    make_placeholders()
    doc = Document()
    setup(doc)
    cover(doc)

    heading(doc, "一、项目背景", 1)
    para(doc, "“狭域巡航 aiPotRol”是在仓储环境智能巡检方向上进一步细化形成的项目。相比传统宽泛的仓储巡检，本项目更关注仓库中容易被忽略、也更适合小车落地的狭窄通道、货架间通道、管廊式仓储区域和高风险设备段。这些区域通常空间有限、视线受遮挡，人工巡查容易遗漏通道堵塞、物体倾倒、人员闯入和环境异常。")
    para(doc, "项目以润和 iCar 智能小车为移动平台，结合 HarmonyOS 上位机 APP、TCP 网络通信、视频回传、麦克纳姆轮控制、环境监测、雷达通道安全判断、视觉异常识别和告警记录展示，形成一套能够现场演示、能够继续扩展的巡检系统。当前项目并不是停留在概念说明阶段，代码中已经实现了较完整的 APP 页面、状态模型、通信命令和消息解析逻辑。")
    para(doc, "本项目的目标可以概括为三个方面：第一，让小车能够在狭窄仓储通道中被稳定连接、查看和接管；第二，让巡检过程以点位、路线、状态和进度的形式被清楚展示；第三，让环境、雷达、视觉和巡检过程中的异常统一进入告警记录，便于现场处理和后续复盘。")
    figure(doc, IMG / "ref_1.png", "图 1 仓储狭域/管廊式巡检场景示意", 6.4)

    heading(doc, "二、项目核心功能与已实现内容", 1)
    para(doc, "项目核心功能围绕“连接小车、控制小车、查看巡检、发现异常、记录告警”展开。APP 首页以“狭域巡航”为入口，显示在线状态、当前点位、下一点位、巡检进度和重点告警，用户可以从首页进入环境感知、SLAM 导航、激光雷达、视觉识别和手动接管等模块。代码中的页面清单已经包括 Index、NetworkSettings、RemoteControl、MecanumWheel、PatrolDashboard、PatrolControl、EnvironmentMonitor、AlertRecords、Module1Environment、Module2Navigation、Module3Lidar、Module4Vision 和 Module5AppMonitor。")
    para(doc, "网络连接功能支持配置小车 IP、TCP 指令端口和视频端口。连接成功后，APP 会初始化消息处理器并进入在线同步状态；连接失败时给出提示；没有真实小车时，也可以跳过连接进入离线预览。这一点对课程演示非常实用，因为它能保证页面和业务流程不完全依赖现场网络。")
    para(doc, "手动接管功能已经覆盖方向按键、摇杆控制、麦克纳姆轮控制和急停。底盘控制协议中包含自由控制、按键控制、四轮独立速度控制、保存图片、开始录制、结束录制、启动巡航和中止巡航等命令。麦克纳姆轮控制是项目亮点之一，它能体现小车在狭窄通道中的横移、原地调整和精细避让能力。")
    para(doc, "自动巡检功能已经在 APP 中形成操作入口，支持手动模式和自动巡检模式切换，并支持开始巡检、暂停巡检、继续巡检、返回起点和停止巡检。巡检路线使用 warehouse_A 和 warehouse_B 两类路线标识，默认点位包括 P1 管廊入口、P2 A 区管线左侧、P3 A 区管线右侧、P4 B 区设备段、P5 高风险设备区和 P6 返回起点。APP 可以显示当前点、下一点、巡检状态和进度，异常发生后也可以将对应点位标记为异常。")
    para(doc, "环境监测模块围绕仓储安全常见指标设计，包含温度、湿度、可燃气体、PM2.5、光照、烟雾和气压等数据。代码中已经定义环境快照、环境历史记录和环境阈值，页面可以显示实时数值和最近历史。雷达模块记录前方距离、左侧距离、右侧距离、通道状态和更新时间，通道状态包括通道畅通、减速通行、通道堵塞和通道占用。视觉模块用于表达人员闯入、物体占道、货物异常和红色瓶体演示等内容。")
    para(doc, "告警模块已经实现统一记录结构，能够保存告警时间、点位、类别、消息、等级和图片地址，并预留 AI 分析结果。告警类别包括环境异常、通道占用、视觉异常和巡检异常。告警记录页支持查看单条告警分析，也支持围绕告警数据进行问答式说明，这部分让项目在异常处置上不只是提示“有问题”，还能进一步给出风险和处理建议。")
    table(
        doc,
        ["模块", "当前实现内容", "展示重点"],
        [
            ("APP 首页", "在线状态、当前点位、下一点位、巡检进度、重点告警、模块入口。", "让评审一眼看到项目完整入口。"),
            ("网络连接", "IP、TCP 端口、视频端口配置，连接成功/失败提示，离线预览。", "说明系统能连接真实小车，也能稳定演示。"),
            ("手动接管", "方向按键、摇杆、麦克纳姆轮、急停、视频辅助观察。", "突出狭窄通道中的可控性和安全性。"),
            ("自动巡检", "手动/自动模式切换，开始、暂停、继续、停止、返回起点。", "形成任务控制流程。"),
            ("环境监测", "温湿度、可燃气体、PM2.5、光照、烟雾、气压与阈值等级。", "体现仓储安全监测。"),
            ("雷达通道", "前方、左侧、右侧距离，通道畅通、减速、堵塞、占用状态。", "体现通道安全能力。"),
            ("视觉识别", "人员闯入、物体占道、红色瓶体演示、蜂鸣器控制。", "体现 AI 巡检亮点。"),
            ("告警记录", "异常类别、等级、点位、时间、图片地址和 AI 分析。", "体现异常复盘能力。"),
        ],
        [3, 8, 5],
    )
    figure(doc, PROTO / "Index.png", "图 2 APP 首页：状态、告警与模块入口", 2.9)
    figure(doc, PROTO / "RemoteControl1.png", "图 3 手动接管页面：视频与遥控控制", 2.9)
    figure(doc, PROTO / "MecanumWheel.png", "图 4 麦克纳姆轮控制：适配狭域横移与原地调整", 2.9)

    heading(doc, "三、应用场景说明", 1)
    para(doc, "项目场景沿用 V3.1 中较清楚的狭域巡检表达，但内容上进一步贴合当前实现。系统的典型路线不是完整大仓库盘点，而是从入口进入管廊或货架通道，依次经过若干固定点位，在重点区域停留观察，并在发现异常后把异常记录到 APP。")
    para(doc, "第一类场景是狭窄通道巡检。小车从入口出发，沿 P1 到 P6 点位移动，APP 展示当前点、下一点和巡检进度。第二类场景是通道占用。用户可以在通道中放置纸箱、瓶体或其他物体，雷达模块显示通道占用或堵塞，视觉模块记录物体异常。第三类场景是人员闯入。人员进入警戒区或巡检通道后，视觉模块产生人员异常记录，APP 告警页展示时间、点位和等级。第四类场景是环境异常。通过真实传感器或模拟数据触发温度、可燃气体、PM2.5、光照等异常，环境模块和告警模块同步展示。第五类场景是人工接管。自动巡检遇到异常时，用户可以切换手动模式，结合视频画面控制小车停止、后退、横移或返回起点。")
    figure(doc, IMG / "ref_3.png", "图 5 点位化巡检路线示意", 4.5)
    figure(doc, IMG / "ref_4.png", "图 6 RViz 建图效果，可用于路线规划和点位标注", 6.2)
    figure(doc, IMG / "ref_5.png", "图 7 巡检点位与异常区域标注", 5.5)

    heading(doc, "四、功能性说明", 1)
    heading(doc, "4.1 上位机 APP 功能", 2)
    para(doc, "APP 是项目最主要的展示和操作入口。首页需要承担总览作用，用户打开应用后应能直接看到项目名称、设备在线状态、当前巡检点位、下一巡检点位、巡检进度和近期重点告警。当前 Index 页面已经实现这些内容，并且通过模块卡片把环境感知、SLAM 导航、激光雷达、视觉识别和手动接管串联起来。")
    para(doc, "网络设置功能用于真实小车联调。用户可以配置默认 IP、指令端口和视频端口，连接成功后进入在线同步状态，连接失败时给出提示。APP 保留离线预览模式，用于在没有真实小车或网络环境不稳定时展示页面结构和业务流程。")
    para(doc, "巡检监控功能用于展示小车自动巡检的实时状态。页面显示路线名称、当前点位、下一点位、巡检状态、巡检进度、雷达状态、视觉状态和最新告警。PatrolDashboard 页面已经按这个思路实现，并使用点位卡片展示路线中的已巡检点和异常点。")
    para(doc, "巡检控制功能用于下发自动巡检命令。用户可以在手动遥控和自动巡检之间切换，可以启动巡检、暂停巡检、继续巡检、停止巡检和返回起点。当前 PatrolControl、Module2Navigation 和 Module5AppMonitor 页面都接入了巡检控制工具，底层通过 JSON 命令向小车端发送 start_patrol、pause_patrol、resume_patrol、return_home 和 stop_patrol。")

    heading(doc, "4.2 小车运动与手动接管功能", 2)
    para(doc, "手动接管是项目安全性的关键。小车在自动巡检过程中如果遇到通道堵塞、人员靠近、识别异常或导航不稳定，用户应能够立即切换到手动模式。当前项目已经实现方向按键、摇杆控制、麦克纳姆轮控制和急停入口，能够覆盖前进、后退、左平移、右平移、左旋转、右旋转、刹车停止等基本动作。")
    para(doc, "麦克纳姆轮控制是狭域巡航场景的重要支撑。普通两轮或四轮小车在货架通道内调头和避让不方便，而麦克纳姆轮可以横向移动，也可以原地调整姿态，更适合狭窄通道、管廊式区域和货架间通道。项目中的 MecanumWheel 页面和四轮独立速度控制协议为这一能力提供了展示入口。")

    heading(doc, "4.3 巡检路线、点位与任务状态功能", 2)
    para(doc, "巡检路线采用点位化方式组织，不追求复杂的仓库全域调度，而是围绕可演示、可定位、可复盘的 P 点路线展开。每个点位都包含是否已巡检和是否异常两个状态。任务状态包括空闲、巡检中、已暂停、正在导航、避障中、异常告警中和返回起点。APP 接收到小车端 patrol 类型消息后，会更新当前点、下一点、进度、路线和状态说明。")
    para(doc, "雷达检测到通道堵塞或占用时，车辆状态可切换为避障中；告警消息指示 alerting 时，车辆状态可切换为异常告警中。这样的状态设计让巡检过程不只是按钮操作，而是能够表现出小车正在做什么、下一步去哪里、当前是否安全。")

    heading(doc, "4.4 环境监测功能", 2)
    para(doc, "环境监测面向仓储安全中的常见隐患。系统展示温度、湿度、可燃气体、PM2.5、光照、烟雾和气压等指标，其中温度、湿度、可燃气体、PM2.5 和光照已经在模型中定义阈值。环境数据不仅显示当前值，还形成历史记录，便于说明异常是在巡检过程中发现的，而不是孤立的一次展示。")
    para(doc, "环境异常采用正常、预警、严重异常三个等级。轻微超限时 APP 可以作为预警显示，严重超限时进入重点告警。环境记录应绑定点位，例如“P5 高风险设备区可燃气体异常”或“P3 A 区管线右侧光照不足”，这样在后续复盘时能够直接定位到异常区域。")

    heading(doc, "4.5 雷达通道安全功能", 2)
    para(doc, "雷达通道安全功能主要用于判断小车前方和左右两侧是否具备安全通行条件。系统记录前方距离、左侧距离和右侧距离，并根据距离判断通道畅通、需要减速、通道堵塞或通道被占用。对于狭窄仓储通道，这一功能可以解释小车为什么需要低速巡检，也可以配合手动接管说明安全策略。")
    para(doc, "当前代码已经定义 LidarSnapshot 和 ChannelStatus，并在消息处理器中支持 lidar 类型消息。收到通道堵塞或通道占用状态后，系统会同步更新雷达状态，并把车辆状态切换为避障中。后续接入真实雷达数据后，可以直接沿用现有页面和状态结构。")

    heading(doc, "4.6 视觉识别与演示功能", 2)
    para(doc, "视觉识别功能用于展示人员闯入、物体占道、货物异常和特定目标检测。当前模型结构已经包含人员检测、箱体或物体进入通道、图像地址和检测说明。Module4Vision 页面还接入了红色瓶体演示相关命令，包括启动演示、停止演示、急停演示、停止蜂鸣器和获取演示状态。")
    para(doc, "红色瓶体演示适合作为答辩中的视觉识别样例。系统可以通过 6600 端口请求演示状态，读取距离阈值、速度、是否检测到目标、蜂鸣器是否激活和更新时间。这个功能说明项目并不只是做了静态页面，而是已经为视觉识别和小车动作联动预留了实际控制链路。")

    heading(doc, "4.7 告警记录与 AI 辅助分析功能", 2)
    para(doc, "告警记录用于把环境、雷达、视觉和巡检过程中的异常统一保存。每条告警包含编号、时间、点位、类别、消息、等级、图片地址和可选 AI 分析结果。首页只展示重点告警，详细记录在 AlertRecords 页面中查看。告警发生后，系统还会尝试触发 AI 分析，生成摘要、风险和处理建议。")
    para(doc, "AI 辅助分析不是项目运行的硬性依赖，但它能增强展示效果。即使分析失败，系统也会给出网络异常提示，不影响告警记录本身。这样的设计保证了基础巡检功能稳定，同时为后续智能化处理留下扩展空间。")

    heading(doc, "五、非功能性说明", 1)
    para(doc, "在易用性方面，系统应尽量让仓库管理员通过 APP 完成主要操作。连接小车、查看状态、启动巡检、暂停巡检、手动接管、查看告警都应在移动端完成，不要求普通用户进入 Linux 终端或 ROS 调试界面。页面文字应接近现场业务表达，例如当前点位、下一点位、通道占用、人员闯入、可燃气体异常等，避免只显示底层技术词。")
    para(doc, "在实时性方面，手动控制指令需要尽量短延迟，尤其是急停、刹车和停止巡检。视频画面应能够辅助用户判断小车前方环境，环境数据、雷达状态和视觉状态应在收到消息后及时刷新。告警记录应优先展示严重异常，避免重要信息被普通日志淹没。")
    para(doc, "在可靠性方面，系统要允许真实连接失败时继续进行页面展示。APP 已经支持离线预览，这是课堂演示和项目答辩中的必要保障。通信层解析 JSON 消息时应容忍分包和残留数据，当前 PatrolMessageHandler 已通过缓冲区方式处理按换行分割的消息。对于未知消息，系统不会直接中断运行。")
    para(doc, "在安全性方面，手动接管、急停和停止巡检的优先级应高于普通巡检动作。小车处于狭窄通道时应优先保证低速和可控，通道堵塞、人员闯入和严重环境异常都应能够触发告警。对于红色瓶体演示和蜂鸣器控制，系统提供停止演示、急停演示和停止蜂鸣器入口，避免演示状态失控。")
    para(doc, "在可维护性方面，项目代码已经按页面、组件、通信、状态、模型和工具类分层。页面位于 pages 目录，视频和通用 UI 位于 components 目录，TCP 连接和巡检命令位于 tcp 目录，状态集中在 PatrolStore，数据结构集中在 PatrolModels。这样的结构便于后续继续接入真实传感器、替换视觉模型、调整页面和扩展任务报告。")
    para(doc, "在可扩展性方面，当前系统已经保留了 SLAM 导航、环境阈值、雷达通道状态、视觉检测日志、AI 告警分析、图片地址和任务状态等接口。后续可以在不大改 APP 页面结构的前提下，接入真实 ROS 节点、YOLO 推理结果、任务报告导出和历史数据查询。")

    heading(doc, "六、概要设计", 1)
    para(doc, "系统整体采用上位机 APP、通信控制层、状态管理层、感知与算法层、小车硬件层的分层设计。上位机 APP 面向用户，负责页面展示和操作入口；通信控制层负责 TCP 连接、命令发送和消息解析；状态管理层负责保存车辆、环境、雷达、视觉、告警和点位数据；感知与算法层负责雷达、视觉、环境阈值和后续 SLAM 导航；硬件层负责底盘运动、传感器采集、视频回传和蜂鸣器等外设控制。")
    para(doc, "APP 端页面结构清晰：首页负责总览，NetworkSettings 负责连接配置，RemoteControl 和 MecanumWheel 负责手动控制，PatrolDashboard 和 PatrolControl 负责巡检监控与任务控制，EnvironmentMonitor 和 Module1Environment 负责环境数据，Module3Lidar 负责通道安全，Module4Vision 负责视觉识别和红色瓶体演示，AlertRecords 负责告警记录和 AI 分析。")
    para(doc, "通信协议分为两类。一类是底盘控制协议，使用帧格式发送自由控制、按键控制、四轮速度、保存图片、视频录制和巡航控制等命令；另一类是巡检业务 JSON 协议，用于发送 start_patrol、pause_patrol、resume_patrol、return_home、stop_patrol、set_mode 等任务指令，并接收 env、patrol、alert、lidar、vision 和 status 等类型的状态消息。")
    para(doc, "状态管理采用集中式 Store。PatrolStore 保存当前环境快照、环境历史、小车状态、雷达状态、视觉状态、视觉日志、巡检点位和告警记录。页面通过订阅 Store 刷新视图，通信层收到消息后更新 Store。这样可以减少页面之间的数据重复，也便于在未来加入本地持久化或报告导出功能。")
    para(doc, "巡检过程可以概括为：用户连接小车，选择手动或自动模式；自动模式下启动巡检路线，小车按点位推进；环境、雷达、视觉模块持续或按点位返回状态；出现异常时写入告警记录并更新点位状态；用户根据视频和告警决定继续巡检、暂停巡检、返回起点或手动接管。")

    heading(doc, "七、截图占位与后续补充", 1)
    para(doc, "V3.1 中用于解释场景和路线的图片继续保留，同时本版为尚未放入正式截图的功能页面预留了图片文件。后续你只需要用最新截图覆盖 generated_assets/project_doc_placeholders_v41 目录下的同名文件，再重新生成文档即可。")
    for idx, (filename, title, _) in enumerate(PLACEHOLDERS, 1):
        figure(doc, PLACEHOLDER_DIR / filename, f"图 8-{idx} {title}：{filename}", 4.8)

    heading(doc, "八、后续完善方向", 1)
    numbered(
        doc,
        [
            "优先补充最终 APP 截图和实物演示图片，尤其是网络连接、巡检监控、雷达通道、视觉识别和告警记录页面。",
            "进一步接入真实环境传感器，减少模拟数据比例，让温湿度、可燃气体、PM2.5 和光照等指标来自真实采集。",
            "完善 SLAM 地图保存、点位配置和自动巡检路线管理，使路线不只依赖预设点位。",
            "补充仓储异常图片数据集，训练更贴合本项目的视觉识别模型，提升人员闯入、物体占道和货物异常的识别稳定性。",
            "增加巡检报告导出功能，把环境数据、异常截图、巡检点位和告警处理建议汇总成单次任务报告。",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("狭域巡航 aiPotRol 项目说明书 V4.1")
    set_font(r, 9, False, "888888")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
