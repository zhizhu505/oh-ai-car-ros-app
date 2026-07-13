from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "仓储狭域巡航与货物异常监测小车-项目需求说明书V3.0.docx"
IMG = ROOT / "generated_assets" / "ref_images"
PROTO = ROOT / "doc" / "prototype"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(9)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], "2F5597")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cells[i], val)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_figure(doc, path, caption, width=6.4):
    path = Path(path)
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(89, 89, 89)


def add_kv_table(doc, pairs):
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(pairs):
        set_cell_text(table.cell(i, 0), k, bold=True)
        set_cell_shading(table.cell(i, 0), "D9EAF7")
        p = table.cell(i, 1).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.cell(i, 1).text = str(v)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.color.rgb = RGBColor(31, 78, 121)
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11)


def add_cover(doc):
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("仓储狭域巡航与货物异常监测小车")
    r.bold = True
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("项目需求说明书 V3.0")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    add_figure(doc, IMG / "ref_1.png", "图 0-1 细分应用场景：仓储狭域通道/管廊式货架巡检闭环", width=6.6)
    add_kv_table(
        doc,
        [
            ("项目名称", "仓储狭域巡航与货物异常监测小车"),
            ("项目定位", "面向仓库狭窄通道、管廊式货架、异常物体占道与人员闯入的移动巡检系统"),
            ("交付形态", "润和 iCar 智能小车 + ROS/SLAM/雷达/视觉能力 + HarmonyOS 上位机 APP"),
            ("文档版本", "V3.0"),
            ("编制日期", "2026 年 7 月 13 日"),
            ("适用阶段", "最终答辩、项目归档、后续迭代开发"),
        ],
    )
    doc.add_page_break()


def add_manual_toc(doc):
    doc.add_heading("目录", level=1)
    entries = [
        "1 项目概述",
        "2 细分应用场景与边界",
        "3 已完成效果与项目亮点",
        "4 用户角色与核心业务流程",
        "5 功能性需求",
        "6 非功能性需求",
        "7 系统架构与接口要求",
        "8 验收标准与答辩展示方案",
        "9 风险、约束与迭代计划",
        "附录 A 术语与修订记录",
    ]
    for e in entries:
        doc.add_paragraph(e, style="List Number")
    doc.add_page_break()


def build_doc():
    doc = Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    add_cover(doc)
    add_manual_toc(doc)

    doc.add_heading("1 项目概述", level=1)
    doc.add_heading("1.1 项目背景", level=2)
    doc.add_paragraph(
        "本项目来源于智能小车实训开发。原始方向是“仓储环境智能巡检与货物异常监测”，"
        "在实际实现过程中，项目组将场景进一步收拢为更适合小车落地演示的狭域巡航场景："
        "仓库货架间窄通道、管廊式仓储通道、设备线缆区、临时堆放区和警戒区。"
        "这些区域空间窄、盲区多、人工巡查频率低，一旦出现货物占道、物品倾倒、人员闯入或环境异常，"
        "管理人员往往无法第一时间发现。"
    )
    doc.add_paragraph(
        "因此，本项目不追求把所有仓储业务一次性做大，而是围绕“能跑、能看、能控、能告警、能复盘”"
        "五件事做成闭环：小车按点位巡航，APP 下发和接管任务，雷达判断通道安全，视觉识别异常物体和人员，"
        "环境数据辅助判定风险，告警记录在移动端集中展示。"
    )

    doc.add_heading("1.2 项目目标", level=2)
    add_bullets(
        doc,
        [
            "替代人工完成狭窄通道和重点货架点位的日常巡检。",
            "通过 HarmonyOS APP 实现连接配置、实时画面、手动遥控、自动巡检和告警查看。",
            "利用麦克纳姆轮底盘实现前进、后退、横移、原地转向等灵活动作，适配狭域通道。",
            "通过 SLAM 地图、巡检点位和通道检测，将异常绑定到具体位置。",
            "将环境、雷达、视觉、任务状态统一整理为可展示、可追溯的巡检记录。",
        ],
    )

    doc.add_heading("1.3 设计原则", level=2)
    add_table(
        doc,
        ["原则", "说明"],
        [
            ("场景收敛", "聚焦仓储狭域通道和管廊式货架，不把系统写成泛化仓库平台。"),
            ("演示可信", "需求描述必须能对应到已完成 APP 页面、控制协议、视频回传和巡检点位。"),
            ("人机协同", "自动巡检不是完全替代人工，异常发生时允许人工一键接管。"),
            ("分层解耦", "APP、通信、运动控制、感知、算法和数据记录分层实现，便于后续替换。"),
            ("渐进交付", "已完成能力写成当前版本交付，SLAM/AI 训练增强等写成明确迭代项。"),
        ],
        widths=[3.5, 12],
    )

    doc.add_heading("2 细分应用场景与边界", level=1)
    doc.add_heading("2.1 场景定义", level=2)
    doc.add_paragraph(
        "本项目的最终场景定义为：面向仓储狭域通道与管廊式货架区域的智能巡检与异常监测。"
        "这里的“货物异常”不仅指传统货架上的货物倾倒，也包括周转箱、瓶体、工具包等物体进入通道、"
        "遮挡巡检路线、遗留在警戒区或出现在不应出现的位置。"
    )
    add_figure(doc, IMG / "ref_3.png", "图 2-1 点位化巡检路线：入口、管廊/货架通道、异常物体、终点", width=4.8)

    doc.add_heading("2.2 典型业务场景", level=2)
    add_table(
        doc,
        ["编号", "场景", "触发方式", "系统表现"],
        [
            ("SC-01", "狭域通道巡检", "APP 选择管廊主线/支线并启动", "小车按 P0-P6 点位巡航，显示当前点、下一点和进度。"),
            ("SC-02", "货物/物体占道", "通道内放置箱体、瓶体或杂物", "雷达提示通道占用，视觉记录异常物体，APP 生成告警。"),
            ("SC-03", "人员闯入", "非巡检人员进入警戒区或通道分支", "视觉识别人员，告警记录绑定点位，必要时人工接管。"),
            ("SC-04", "环境异常", "模拟温湿度、烟雾、可燃气体、光照异常", "APP 环境卡片变为预警/严重状态，记录异常时间和点位。"),
            ("SC-05", "手动接管", "自动巡检受阻或演示需要", "切换手动模式，摇杆/按键控制小车，视频画面辅助判断。"),
            ("SC-06", "远程联调", "网络连接真实小车或离线预览", "支持 IP、TCP 端口、视频端口配置；离线可预览全部页面。"),
        ],
        widths=[2, 3, 4, 7],
    )

    doc.add_heading("2.3 场景边界", level=2)
    add_table(
        doc,
        ["范围", "纳入", "暂不纳入"],
        [
            ("空间范围", "货架间通道、管廊式通道、警戒区、充电/起终点", "大型叉车调度、全仓库存盘点。"),
            ("异常范围", "通道占用、人员闯入、货物/物体倾倒、环境指标异常", "复杂商品 SKU 识别、库存数量自动核算。"),
            ("控制范围", "APP 连接、手动遥控、自动巡检、暂停/继续/停止/返航", "多车协同调度、云端跨仓库管理。"),
            ("数据范围", "实时状态、环境历史、视觉日志、告警记录、点位进度", "长期大数据分析平台。"),
        ],
        widths=[3, 6, 6],
    )

    doc.add_heading("3 已完成效果与项目亮点", level=1)
    doc.add_heading("3.1 当前版本已完成能力", level=2)
    add_table(
        doc,
        ["模块", "已完成效果", "答辩展示价值"],
        [
            ("HarmonyOS APP", "首页、巡检总览、巡检控制、环境监测、告警记录、五个模块页、网络连接页已成型。", "能直接展示项目完整交互入口，不只是代码接口。"),
            ("小车通信", "支持 TCP 指令端口 6000、视频端口 6500；可配置 IP 和端口；连接后接收并解析状态。", "可以现场说明真实小车连接链路。"),
            ("远程控制", "支持摇杆、方向按键、麦克纳姆轮控制、急停、速度控制等指令封装。", "体现小车运动控制和狭域机动能力。"),
            ("自动巡检控制", "支持手动/自动模式切换，开始、暂停、继续、停止、返回起点等巡检指令。", "形成“任务下发-执行-接管”的业务闭环。"),
            ("状态看板", "展示在线状态、当前点位、下一点位、巡检进度、通道状态、视觉状态、最新告警。", "管理人员不需要看终端日志即可理解现场状态。"),
            ("环境与告警", "环境数据可历史化展示，告警按等级、类别、点位记录；支持 AI 辅助分析字段。", "突出异常处理和复盘能力。"),
            ("视频回传", "APP 集成视频组件，可按配置的视频端口展示小车视角。", "提升现场演示直观性。"),
        ],
        widths=[3, 8, 5],
    )

    doc.add_heading("3.2 亮点总结", level=2)
    add_bullets(
        doc,
        [
            "场景更具体：从“大仓库巡检”收敛为“狭域通道/管廊式货架巡检”，更符合小车体积、麦轮底盘和课堂演示条件。",
            "控制链路完整：APP 能连接小车、发送控制指令、展示视频画面，并支持断网或未连接时的离线预览。",
            "点位感明显：P0-P6 巡检点把路线、当前进度、异常位置和演示脚本串起来，讲解成本低。",
            "异常类型可演示：通道堵塞、人员闯入、环境异常、货物/物体异常都能通过实物或模拟数据展示。",
            "人工可接管：自动巡检出现异常时，管理员能快速切回手动模式，符合真实仓储安全要求。",
            "后续扩展清楚：SLAM 建图、YOLO 训练、数据归档、低电返航等都有明确接口和验收方向。",
        ],
    )
    add_figure(doc, PROTO / "Index.png", "图 3-1 APP 首页：状态、告警、五大模块入口", width=3.0)
    add_figure(doc, PROTO / "RemoteControl1.png", "图 3-2 手动接管页面：视频与控制入口", width=3.0)
    add_figure(doc, PROTO / "MecanumWheel.png", "图 3-3 麦克纳姆轮控制：适配狭域横移与原地调整", width=3.0)

    doc.add_heading("4 用户角色与核心业务流程", level=1)
    doc.add_heading("4.1 用户角色", level=2)
    add_table(
        doc,
        ["角色", "关注点", "主要操作"],
        [
            ("仓库管理员", "是否有异常、异常在哪里、是否需要处理", "启动巡检、查看告警、查看实时画面、手动接管。"),
            ("巡检运维人员", "设备是否在线、通信是否正常、地图和点位是否可靠", "配置 IP/端口、调试底盘、查看日志、维护传感器。"),
            ("项目展示人员", "能否清晰讲出项目亮点并完成演示", "按演示脚本触发通道堵塞、环境异常、人员闯入等场景。"),
            ("后续开发人员", "模块接口是否清楚、需求是否可继续迭代", "扩展 YOLO 模型、完善 ROS 节点、增强数据归档。"),
        ],
        widths=[3, 6, 7],
    )

    doc.add_heading("4.2 核心业务流程", level=2)
    add_bullets(
        doc,
        [
            "准备阶段：小车上电，APP 输入小车 IP、TCP 端口和视频端口，建立连接。",
            "任务阶段：管理员选择路线并启动自动巡检，系统显示当前点位、下一点位和巡检进度。",
            "感知阶段：到达点位后采集环境、雷达和视觉状态，发现异常则记录点位和类别。",
            "处置阶段：APP 弹出或记录告警；严重异常时可以暂停巡检或切换手动遥控。",
            "复盘阶段：在告警记录、环境历史和视觉日志中查看异常过程，为答辩或后续维护提供依据。",
        ],
    )
    add_figure(doc, IMG / "ref_5.png", "图 4-1 基于 RViz 地图的点位路线与异常标注示意", width=5.7)

    doc.add_heading("5 功能性需求", level=1)
    doc.add_paragraph("需求编号采用 FR-模块-序号。优先级 P0 为当前版本必须稳定展示，P1 为应完成或已具备基础，P2 为后续增强。")

    doc.add_heading("5.1 APP 连接与监控模块（FR-APP）", level=2)
    add_table(
        doc,
        ["编号", "优先级", "需求描述", "验收标准"],
        [
            ("FR-APP-01", "P0", "支持输入小车 IP、TCP 端口、视频端口并发起连接。", "连接成功后首页显示在线；连接失败有提示；可跳过进入离线预览。"),
            ("FR-APP-02", "P0", "首页显示在线状态、当前点位、下一点位、巡检进度和重点告警。", "状态变化能在页面刷新，告警按严重程度优先展示。"),
            ("FR-APP-03", "P0", "巡检控制页支持手动/自动模式切换。", "点击后本地状态更新，并向小车端发送对应控制命令。"),
            ("FR-APP-04", "P0", "支持开始、暂停、继续、停止、返回起点等巡检操作。", "命令发送后有明确反馈，断网时保留本地状态提示。"),
            ("FR-APP-05", "P0", "支持视频画面显示，用于远程观察小车前方环境。", "视频端口可配置，页面布局不遮挡控制区域。"),
            ("FR-APP-06", "P1", "告警记录页按类别、等级、点位展示异常。", "环境、雷达、视觉类告警均能进入记录列表。"),
        ],
        widths=[2.5, 1.8, 7, 5],
    )

    doc.add_heading("5.2 运动控制与小车接管模块（FR-MOVE）", level=2)
    add_table(
        doc,
        ["编号", "优先级", "需求描述", "验收标准"],
        [
            ("FR-MOVE-01", "P0", "支持前进、后退、左平移、右平移、左旋转、右旋转、刹车停止。", "按键控制响应稳定，急停优先级最高。"),
            ("FR-MOVE-02", "P0", "支持摇杆无级控制，并将 X/Y 方向编码为小车控制指令。", "摇杆松手后小车停止，方向与页面反馈一致。"),
            ("FR-MOVE-03", "P0", "支持麦克纳姆轮四轮速度控制。", "能展示横移和原地调整动作，适合狭域通道纠偏。"),
            ("FR-MOVE-04", "P1", "支持速度档位或速度比例调节。", "低速档用于窄通道巡检，高速档用于空旷区域转场。"),
        ],
        widths=[2.5, 1.8, 7, 5],
    )

    doc.add_heading("5.3 巡检路线与点位模块（FR-PATROL）", level=2)
    add_table(
        doc,
        ["编号", "优先级", "需求描述", "验收标准"],
        [
            ("FR-PATROL-01", "P0", "系统内置管廊主线和管廊支线两类路线。", "APP 可显示当前路线，巡检控制命令携带路线标识。"),
            ("FR-PATROL-02", "P0", "巡检点采用 P0-P6 编号，点位名称可读。", "总览页展示点位卡片，已巡检和异常点有不同状态。"),
            ("FR-PATROL-03", "P1", "自动巡检时按点位推进，支持当前点和下一点展示。", "模拟或真实数据进入后，进度条和点位状态同步变化。"),
            ("FR-PATROL-04", "P1", "支持异常点标记。", "告警发生后，对应点位在路线中标为异常。"),
        ],
        widths=[2.5, 1.8, 7, 5],
    )

    doc.add_heading("5.4 环境感知模块（FR-ENV）", level=2)
    add_table(
        doc,
        ["编号", "优先级", "需求描述", "验收标准"],
        [
            ("FR-ENV-01", "P0", "APP 展示温度、湿度、烟雾/可燃气体、光照、PM2.5 等指标。", "环境页面能刷新数值，并保留最近历史记录。"),
            ("FR-ENV-02", "P1", "支持按仓储安全阈值判定正常、预警、严重三类状态。", "超过阈值后产生环境告警，包含时间、点位、指标和值。"),
            ("FR-ENV-03", "P1", "环境异常与巡检点位绑定。", "记录可读，例如“P3 窄通道可燃气体异常”。"),
            ("FR-ENV-04", "P2", "支持阈值配置持久化。", "APP 重启后阈值仍生效。"),
        ],
        widths=[2.5, 1.8, 7, 5],
    )

    doc.add_heading("5.5 雷达通道安全模块（FR-LIDAR）", level=2)
    add_table(
        doc,
        ["编号", "优先级", "需求描述", "验收标准"],
        [
            ("FR-LIDAR-01", "P1", "利用雷达判断前方、左侧、右侧安全距离。", "通道过窄或障碍过近时 APP 显示通道异常。"),
            ("FR-LIDAR-02", "P1", "通道被箱体、工具包等大物体占用时生成堵塞告警。", "告警包含点位、占用描述和建议处置方式。"),
            ("FR-LIDAR-03", "P1", "自动巡检中遇到近距离障碍时减速或停止。", "0.4m 内提示障碍，0.2m 内触发停止策略。"),
            ("FR-LIDAR-04", "P2", "与 SLAM/导航节点联动实现绕障或重新规划。", "临时障碍移除后可继续巡检。"),
        ],
        widths=[2.5, 1.8, 7, 5],
    )

    doc.add_heading("5.6 视觉识别与异常记录模块（FR-VISION）", level=2)
    add_table(
        doc,
        ["编号", "优先级", "需求描述", "验收标准"],
        [
            ("FR-VISION-01", "P1", "识别人员闯入、通道内箱体/瓶体/异常物体。", "画面或日志能区分人员与物体异常。"),
            ("FR-VISION-02", "P1", "货物异常使用可演示物体模拟，包括倾倒、占道、缺失或越界。", "异常发生后生成视觉日志和告警记录。"),
            ("FR-VISION-03", "P1", "视觉异常与巡检点位绑定。", "记录中包含 P 点编号、异常类别、时间和说明。"),
            ("FR-VISION-04", "P2", "后续接入 YOLOv5/YOLOv8 训练模型并进行 TensorRT 加速。", "在 Jetson 端实现实时检测，画面不卡顿。"),
            ("FR-VISION-05", "P2", "支持目标跟随或指定目标观察。", "APP 选择目标后，小车保持安全距离跟随或观察。"),
        ],
        widths=[2.5, 1.8, 7, 5],
    )

    doc.add_heading("5.7 数据记录与告警分析模块（FR-DATA）", level=2)
    add_table(
        doc,
        ["编号", "优先级", "需求描述", "验收标准"],
        [
            ("FR-DATA-01", "P0", "统一维护车辆、环境、雷达、视觉、点位和告警状态。", "页面切换后状态不丢失。"),
            ("FR-DATA-02", "P0", "告警记录支持等级、类别、消息、点位、图片地址等字段。", "首页和告警页读取同一份记录。"),
            ("FR-DATA-03", "P1", "环境历史保留最近数据，用于曲线或列表展示。", "连续刷新时历史记录长度受控，不影响页面流畅度。"),
            ("FR-DATA-04", "P1", "支持 AI 辅助分析告警，给出摘要、风险和处理建议。", "分析失败时给出网络异常提示，不影响告警记录。"),
            ("FR-DATA-05", "P2", "后续将任务报告、截图、轨迹和环境日志落盘保存。", "单次巡检可导出报告。"),
        ],
        widths=[2.5, 1.8, 7, 5],
    )

    doc.add_heading("6 非功能性需求", level=1)
    add_table(
        doc,
        ["类别", "编号", "指标要求"],
        [
            ("性能", "NFR-PERF-01", "手动控制指令端到端响应目标不高于 200ms，急停优先。"),
            ("性能", "NFR-PERF-02", "APP 页面切换、告警列表和环境数据刷新保持流畅，无明显卡顿。"),
            ("可靠性", "NFR-REL-01", "网络断开时 APP 给出明确提示，不误显示为正常在线。"),
            ("可靠性", "NFR-REL-02", "自动巡检受阻时可暂停、停止或切换手动接管。"),
            ("可用性", "NFR-USE-01", "管理员不需要 Linux 终端即可完成连接、查看、控制和告警确认。"),
            ("可维护性", "NFR-MAIN-01", "APP 页面、通信工具、状态存储、模型定义分层清晰，便于后续扩展。"),
            ("安全性", "NFR-SAFE-01", "急停和停止巡检始终高于普通运动命令。"),
            ("安全性", "NFR-SAFE-02", "低速通过窄通道，通道堵塞和人员闯入按严重告警处理。"),
            ("兼容性", "NFR-COMP-01", "支持真实小车连接和离线预览两种运行方式，便于课堂展示。"),
        ],
        widths=[2.5, 3, 10],
    )

    doc.add_heading("7 系统架构与接口要求", level=1)
    doc.add_heading("7.1 总体架构", level=2)
    add_table(
        doc,
        ["层级", "组成", "职责"],
        [
            ("终端交互层", "HarmonyOS / ArkTS / ArkUI APP", "连接配置、视频显示、巡检控制、环境与告警展示。"),
            ("通信网关层", "TCP 长连接、视频流、消息解析", "将 APP 指令转为小车控制消息，将小车状态推送到 APP。"),
            ("业务调度层", "巡检路线、点位状态、异常聚合", "维护巡检任务状态，协调环境、雷达、视觉等信息。"),
            ("感知算法层", "SLAM、雷达避障、视觉识别、环境阈值", "完成地图、定位、通道安全和异常识别。"),
            ("硬件执行层", "iCar 底盘、麦轮、雷达、摄像头、Jetson/MCU", "负责移动、采集、计算和本地运行。"),
        ],
        widths=[3, 5, 8],
    )
    add_figure(doc, IMG / "ref_4.png", "图 7-1 实际 RViz 建图效果：用于路线规划与点位标注", width=6.4)

    doc.add_heading("7.2 通信接口", level=2)
    add_table(
        doc,
        ["接口", "端口/格式", "说明"],
        [
            ("指令通信", "TCP 6000", "APP 与小车端建立长连接，发送运动、巡检、模式切换等指令。"),
            ("视频回传", "TCP/HTTP 视频流 6500", "APP 视频组件展示小车实时视角。"),
            ("基础帧格式", "$ + 车辆类型 + 命令标记 + 数据长度 + 数据主体 + 校验 + #", "用于底盘控制类指令封装。"),
            ("巡检业务消息", "JSON 或约定字段", "用于环境、雷达、视觉、告警和任务状态同步。"),
        ],
        widths=[3, 4, 9],
    )
    add_table(
        doc,
        ["命令", "用途", "当前应用"],
        [
            ("cmd 10", "自由控制小车", "摇杆控制 X/Y 方向。"),
            ("cmd 15", "按键控制小车", "前进、后退、平移、旋转、刹车。"),
            ("cmd 21", "四轮单独更新速度", "麦克纳姆轮独立速度控制。"),
            ("cmd 60", "保存单张图片", "巡检取证或异常抓拍。"),
            ("cmd 61/62", "开始/结束录制视频", "现场视频留存。"),
            ("cmd 63/64", "启动/中止巡航", "自动巡检任务控制。"),
        ],
        widths=[3, 5, 8],
    )

    doc.add_heading("8 验收标准与答辩展示方案", level=1)
    doc.add_heading("8.1 总体验收标准", level=2)
    add_table(
        doc,
        ["验收项", "通过标准"],
        [
            ("APP 完整性", "能从首页进入各模块，页面状态清晰，无明显布局错乱。"),
            ("真实连接", "能配置 IP/端口并尝试连接小车；成功后状态更新。"),
            ("手动控制", "能控制小车完成前后、横移、旋转和停止。"),
            ("自动巡检", "能下发开始/暂停/继续/停止/返航等命令，并展示点位进度。"),
            ("视频展示", "能在 APP 页面展示或预留小车实时画面区域。"),
            ("异常演示", "至少完成通道占用、人员闯入、环境异常三类告警展示。"),
            ("文档一致性", "文档中的已完成效果能在代码或演示视频中找到对应。"),
        ],
        widths=[4, 12],
    )

    doc.add_heading("8.2 推荐答辩演示脚本", level=2)
    add_table(
        doc,
        ["步骤", "展示内容", "讲解重点"],
        [
            ("1", "打开 APP 首页，说明“狭域巡航”定位。", "先讲场景收敛：不是泛仓储，而是狭窄通道和管廊式货架。"),
            ("2", "进入连接弹窗，配置 IP、TCP 6000、视频 6500。", "说明 APP 到小车的真实通信链路。"),
            ("3", "展示巡检总览页。", "说明 P0-P6 点位、当前点、下一点、进度和路线。"),
            ("4", "启动自动巡检，再暂停/继续。", "体现任务控制闭环。"),
            ("5", "放置箱体或瓶体模拟货物异常/通道占用。", "雷达和视觉共同支撑异常判断。"),
            ("6", "人员进入警戒区。", "展示人员闯入告警和点位绑定。"),
            ("7", "模拟温湿度/烟雾/光照异常。", "展示环境数据和分级告警。"),
            ("8", "切换手动接管，使用摇杆或麦轮控制绕开障碍。", "强调人机协同和安全接管。"),
        ],
        widths=[1.6, 6, 8],
    )

    doc.add_heading("8.3 演示素材建议", level=2)
    add_bullets(
        doc,
        [
            "一张路线图：说明 P0-P6 巡检点和异常区域。",
            "一张 RViz 建图截图：证明 SLAM/地图不是纯概念。",
            "一段 APP 操作录屏：展示连接、总览、控制、告警记录。",
            "两个实物异常道具：纸箱/瓶体模拟通道占用，人员进入模拟安防告警。",
            "一组环境异常模拟值：温度高、可燃气体异常、光照不足或 PM2.5 异常。",
        ],
    )

    doc.add_heading("9 风险、约束与迭代计划", level=1)
    add_table(
        doc,
        ["风险/约束", "影响", "处理方式"],
        [
            ("真实小车网络不稳定", "APP 控制或视频可能中断", "保留离线预览；演示前固定热点、IP 和端口。"),
            ("SLAM 地图质量受场地影响", "路线点位可能漂移", "答辩前保存稳定地图，必要时采用固定点位演示。"),
            ("视觉模型训练数据不足", "复杂货物异常识别准确率有限", "使用可控道具演示，后续补充数据集训练。"),
            ("环境传感器接入不完整", "部分环境数据需要模拟", "APP 层保留真实数据接口，演示阶段用模拟数据验证流程。"),
            ("狭窄通道安全风险", "小车可能碰撞货架或障碍", "限制速度，保留急停和手动接管。"),
        ],
        widths=[4.5, 5, 6.5],
    )

    doc.add_heading("9.2 后续迭代计划", level=2)
    add_table(
        doc,
        ["阶段", "目标", "交付物"],
        [
            ("V3.1", "完善真实传感器数据接入和环境阈值配置", "环境数据协议、阈值设置页、真实告警记录。"),
            ("V3.2", "接入 YOLO 模型并完成异常物体/人员实时识别", "标注数据集、推理节点、视觉日志截图。"),
            ("V3.3", "强化 SLAM 导航和断点续巡", "稳定地图、点位配置、任务状态持久化。"),
            ("V3.4", "生成单次巡检报告并支持导出", "任务报告、异常图片、环境曲线、路线轨迹。"),
        ],
        widths=[2.5, 7.5, 6],
    )

    doc.add_heading("附录 A 术语与修订记录", level=1)
    add_table(
        doc,
        ["术语", "说明"],
        [
            ("狭域巡航", "面向狭窄通道、管廊、货架间通道的低速点位巡检。"),
            ("货物异常", "本文中包括货物倾倒、物体占道、遗留物、货架区域缺失或越界。"),
            ("SLAM", "同步定位与建图，用于构建巡检地图并确定小车位置。"),
            ("麦克纳姆轮", "支持横移和原地转向的全向轮结构，适合狭小空间调姿。"),
            ("TCP 6000/6500", "本项目 APP 与小车通信使用的指令端口和视频端口。"),
        ],
        widths=[4, 12],
    )
    add_table(
        doc,
        ["版本", "日期", "修订内容"],
        [
            ("V1.0", "2026-07-08", "形成原始仓储巡检需求。"),
            ("V2.1", "2026-07-10", "补充系统设计、功能编号和概要架构。"),
            ("V3.0", "2026-07-13", "结合实际 APP 和演示效果，收敛为仓储狭域通道/管廊式巡检场景，重写项目需求说明书。"),
        ],
        widths=[2.5, 3.5, 10],
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "仓储狭域巡航与货物异常监测小车 项目需求说明书 V3.0"

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
