from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"E:\test_java\huawei\项目开发及管理文档_狭域巡航_aiPotRol.docx"
LOGO = r"C:\Users\lenovo\AppData\Local\Temp\codex-clipboard-0fbbb741-4bb2-47e2-ac65-3ac3e6ebe386.png"
FONT = "宋体"


def font_run(run, size=10.5, bold=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    font_run(r, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def paragraph(doc, text="", first_line=True, size=10.5, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    font_run(r, size=size, bold=bold)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    font_run(r, size=16 if level == 1 else 13, bold=True)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(t.rows[0].cells[i], "D9EAF7")
        if widths:
            t.rows[0].cells[i].width = widths[i]
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=10, align=WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return t


def gantt(doc):
    dates = ["7.6", "7.7", "7.8", "7.9", "7.10", "7.11", "7.12", "7.13", "7.14", "7.15"]
    tasks = [
        ("企业授课、需求理解、方向确定", [0]),
        ("开发环境与小车连接", [0, 1]),
        ("小车硬件熟悉与基础控制", [1, 2]),
        ("仓储场景规划与布局搭建", [3, 4, 5, 6, 7, 8]),
        ("环境感知代码与数据展示", [2, 3, 6, 7]),
        ("SLAM 建图、路线规划与定点巡检", [2, 3, 6, 7, 8]),
        ("雷达追踪、避障与通道堵塞检测", [2, 3, 6, 7, 8]),
        ("YOLO 数据集、训练与异常识别", [2, 3, 6, 7]),
        ("鸿蒙 APP、通信、看板与告警", [3, 4, 5, 6, 7, 8]),
        ("文档、PPT、中期与结题材料", [3, 4, 5, 6, 7, 8, 9]),
        ("中期答辩", [4]),
        ("集中实车联调与最终演示", [6, 7, 8]),
        ("最终答辩", [9]),
    ]
    t = doc.add_table(rows=1, cols=1 + len(dates))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(t.rows[0].cells[0], "任务", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(t.rows[0].cells[0], "D9EAF7")
    for i, d in enumerate(dates, 1):
        set_cell_text(t.rows[0].cells[i], d, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(t.rows[0].cells[i], "D9EAF7")
    for task, active in tasks:
        cells = t.add_row().cells
        set_cell_text(cells[0], task, size=9)
        for i in range(len(dates)):
            set_cell_text(cells[i + 1], "■" if i in active else "", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            if i in active:
                set_cell_shading(cells[i + 1], "70AD47")
            elif i in [3, 4, 5]:
                set_cell_shading(cells[i + 1], "FCE4D6")
    paragraph(doc, "说明：绿色表示实际投入开发或交付的主要时段，浅橙色表示受小车电池故障、连续暴雨和设备不可用影响，需要调整计划的时段。7.12 至 7.14 为集中补强阶段，团队将开发时间延长至 9:00-21:00，用于补齐实车联调和最终展示能力。")


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.0)

    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    doc.styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目开发及管理文档")
    font_run(r, 22, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("狭域巡航 aiPotRol")
    font_run(r, 18, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("仓储环境智能巡检与货物异常监测小车")
    font_run(r, 15, False)

    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(LOGO, width=Cm(5.2))
    except Exception:
        paragraph(doc, "（Logo 图片未能嵌入，生成时请检查图片路径。）", first_line=False)

    heading(doc, "一、项目基本信息")
    paragraph(doc, "本项目品牌为“狭域巡航 aiPotRol”，项目名称为“仓储环境智能巡检与货物异常监测小车”。项目周期为 2026 年 7 月 6 日至 2026 年 7 月 15 日，最终答辩安排在 2026 年 7 月 15 日全天线下进行。团队成员包括武芷竹、胡雅欣、周诗晴、曹玥。")
    paragraph(doc, "项目面向仓储货架区、狭窄通道、货物堆放区、警戒区和固定巡检点位，核心硬件包括智能小车、激光雷达、视觉相机和环境感知传感器；软件环境包括 Ubuntu、ROS、YOLO 目标检测、HarmonyOS / ArkTS / ArkUI 和 TCP 通信。主要交付包括鸿蒙 APP、仓储巡检演示场景、实车巡检与避障能力、货物/红瓶异常识别能力、项目文档、测试报告、使用手册、答辩 PPT 和演示视频。")

    heading(doc, "二、项目背景与开发目标")
    paragraph(doc, "2026 年 7 月 6 日企业老师授课后，团队明确本次实训需要从单项实验转向完整项目交付，并在 7 月 15 日最终答辩中展示项目立项、开发过程、成员分工、设计实现和实际成果。")
    paragraph(doc, "结合课程要求、硬件条件和团队创意，项目确定为“狭域巡航 aiPotRol：仓储环境智能巡检与货物异常监测小车”。项目以智能小车为移动平台，围绕仓储货架、狭窄通道和货物堆放场景，实现环境感知、SLAM 巡检、雷达避障、视觉异常识别和鸿蒙 APP 远程监控。")

    heading(doc, "三、需求分析")
    heading(doc, "3.1 功能需求", 2)
    table(doc, ["序号", "需求项", "实际落地方式"], [
            ["1", "仓储环境数据采集与异常告警", "采集温湿度、PM2.5、光照、GPS 等数据，并在 APP 中展示状态和异常。"],
            ["2", "SLAM 建图与固定点巡检", "基于仓储场景生成地图，设计货架/通道巡检点和巡检路线。"],
            ["3", "激光雷达避障与通道堵塞检测", "识别通道障碍，支持减速、停止和通道占用提示。"],
            ["4", "视觉识别人员闯入、红瓶/货物异常等目标", "通过 YOLO 训练与部署，完成红瓶停止、警戒区和货物异常识别。"],
            ["5", "鸿蒙 APP 状态展示、数据展示、告警和控制", "提供网络连接、遥控、巡检看板、环境/雷达/视觉页面、告警记录和 AI 咨询。"],
    ])
    heading(doc, "3.2 非功能需求", 2)
    paragraph(doc, "项目需要满足演示稳定、页面清晰、现场可控和过程可追溯要求。考虑到硬件可用时间受限、天气影响取还设备、小车电池和串口问题等风险，团队采用“硬件可用时优先实车联调，硬件不可用时推进 APP、模型、文档、PPT 和场景搭建”的策略，保证总进度不断档。")

    heading(doc, "四、总体设计")
    paragraph(doc, "系统由智能小车硬件、ROS 控制与导航、环境/雷达/视觉感知、TCP 通信、鸿蒙 APP 和仓储演示场景组成。小车负责运动执行，ROS 负责建图导航与控制调度，传感器、雷达和视觉模块负责异常判断，APP 负责连接、展示、控制和告警汇总。")
    table(doc, ["层级", "组成", "作用"], [
        ["硬件执行层", "智能小车、底盘、电池、雷达、相机、传感器", "完成移动、采集、识别和现场演示所需的基础硬件能力。"],
        ["ROS 与算法层", "SLAM、导航、雷达处理、视觉识别、传感器数据处理", "完成建图、巡检路线执行、避障、异常判断和数据输出。"],
        ["通信层", "TCP 控制通道、数据消息处理", "连接鸿蒙 APP 与小车端，实现控制命令发送和状态信息接收。"],
        ["APP 展示层", "网络配置、巡检看板、遥控、环境、雷达、视觉、告警和 AI 咨询页面", "为现场操作和答辩展示提供统一入口。"],
        ["交付层", "仓储场景、PPT、文档、测试报告、使用手册、演示视频", "支撑最终答辩和课程归档。"],
    ])

    heading(doc, "五、项目管理主线与时间推进")
    paragraph(doc, "本项目的核心管理难点在于开发周期短、硬件依赖强、天气和设备状态不可控。团队实际从 7 月 6 日开始正式实践，到 7 月 15 日全天答辩，完整开发窗口不足十天。因此项目管理重点不是简单列出模块，而是每天根据设备状态和答辩节点动态调整工作重心。")
    heading(doc, "5.1 项目进度甘特图", 2)
    gantt(doc)

    heading(doc, "5.2 每日开发与管理记录", 2)
    table(doc, ["日期", "项目状态", "当天重点工作", "管理调整与结果"], [
        ["7.6 周一", "正式启动", "企业老师授课，明确课程要求、答辩目标和智能小车项目方向；完成 Ubuntu、ROS 虚拟机环境搭建。", "建立协作方式，将任务拆分为环境感知、SLAM、雷达、视觉和 APP 五条线。"],
        ["7.7 周二", "硬件上手", "下午拿到小车，熟悉硬件功能，完成实机环境配置、电脑连接和 APP 命令控制验证。", "优先打通小车、虚拟机和 APP 控制链路，为后续开发建立基础。"],
        ["7.8 周三", "多模块并行", "推进传感器代码、实车 SLAM 建图、目标检测数据集、雷达追踪避障和 APP 内容设计。", "进入并行开发阶段，各成员产出可验证结果，并同步准备中期材料。"],
        ["7.9 周四", "硬件受阻", "小车电池无法充电并返工修理，实车调试受限；团队转向传感器代码、YOLO 训练、SLAM 方案、APP 调试、文档和中期 PPT。", "把不依赖小车的工作前置完成，减少硬件等待造成的进度损失。"],
        ["7.10 周五", "中期节点", "连续暴雨影响设备取用，完成中期答辩；继续修改 PPT，整理前期成果、后续规划和鸿蒙 APP 代码。", "借中期答辩重新确认剩余重点：实车联调、APP 展示和雷达/视觉/导航闭环。"],
        ["7.11 周六", "持续受限", "暴雨影响仍在，无法稳定拿到小车；继续完善 APP、文档、演示方案、仓储场景和联调计划。", "将受阻时间转化为材料和软件完善时间，为恢复硬件后快速联调做准备。"],
        ["7.12 周日", "集中补强", "开发时间延长至 9:00-21:00，推进小车控制、视频修复、界面整合、地图参数、定点巡检、AI 接入、横竖屏适配和视觉环境。", "进入冲刺阶段，把分散模块串成可演示流程。"],
        ["7.13 周一", "联调深化", "继续 9:00-21:00 开发，优化 APP 仓储场景页面，完善数据面板、多档调速、红瓶停止、警戒区识别、AI 咨询、雷达接入、地图路线和串口排查。", "从功能可用转向展示完整度，重点解决界面、数据、控制、雷达、视觉和导航联动。"],
        ["7.14 周二", "答辩前收口", "上午小车出现黑屏并再次无法充电，团队将小车送去维修；等待期间继续整理结题 PPT、完善文档和最终演示流程。", "根据硬件状态再次调整节奏，把无法实车调试的时间转为答辩材料收口和演示逻辑梳理。"],
        ["7.15 周三", "最终答辩", "全天线下结题答辩，围绕项目立项、开发过程、成员贡献、设计实现、测试结果和最终成果进行展示。", "以项目过程管理和实际成果为汇报主线，突出短周期、强约束条件下的计划调整和团队协作。"],
    ])

    heading(doc, "5.3 进度偏差与应对")
    paragraph(doc, "原计划中，7 月 7 日至 7 月 11 日应是硬件持续联调阶段，但 7 月 9 日小车电池无法充电、7 月 10 日和 7 月 11 日连续暴雨，导致设备取用和实车调试被压缩。7 月 14 日上午小车又出现黑屏并再次无法充电，团队将小车送去维修，进一步增加了答辩前硬件验证的不确定性。")
    paragraph(doc, "团队将工作拆为“依赖小车”和“不依赖小车”两类：硬件不可用时推进 APP、模型、文档、PPT 和仓储场景；硬件恢复后集中完成小车控制、SLAM、雷达、视觉和 APP 联调。7 月 11 日至 7 月 14 日开发时间延长至 9:00-21:00，其中 7 月 14 日维修等待时间主要用于结题 PPT 整理、文档收口和演示流程梳理。")

    heading(doc, "六、成员分工与贡献")
    table(doc, ["成员", "角色定位", "主要贡献"], [
        ["武芷竹", "组长 / SLAM 导航 / 项目统筹", "负责进度协调、SLAM 建图导航、仓储路线设计、地图参数调整、定点巡检测试、串口排查和联调推进。"],
        ["胡雅欣", "鸿蒙 APP / 环境感知 / 展示材料", "负责 APP 页面与小车控制适配、环境传感器代码、数据面板、多档调速、需求文档和结题 PPT。"],
        ["周诗晴", "雷达避障 / AI 咨询 / APP 功能接入", "负责雷达追踪避障、雷达功能接入 APP、AI 咨询、横竖屏适配、logo 和界面完善。"],
        ["曹玥", "视觉识别 / YOLO 模型 / 场景代码", "负责目标检测数据集、YOLO 模型训练、红瓶停止、警戒区识别、场景行动代码和视觉环境配置。"],
    ])
    paragraph(doc, "分工并非固定不变。中期前团队按照模块划分任务，中期后根据硬件故障、暴雨停工和最终答辩压力进行动态调整：硬件恢复时集中支持实车联调，硬件不可用时共同推进文档、PPT、APP 和场景布置，保证每个时间段都有可交付成果。")

    heading(doc, "七、模块设计与实现概况")
    table(doc, ["模块", "开发内容", "与时间进度的关系"], [
        ["环境感知模块", "温湿度、PM2.5、光照、GPS 数据构建与 APP 图表展示。", "7.8 启动，7.13 完善数据面板。"],
        ["SLAM 导航模块", "实车建图、自动导航、巡检点、路线规划和窄道通行参数。", "7.8-7.9 初步验证，7.12-7.14 集中联调。"],
        ["雷达避障模块", "雷达追踪、自动避障、通道堵塞检测和串口排查。", "7.8 配置，7.12-7.13 与 APP 和底盘联调。"],
        ["视觉识别模块", "数据集构建、YOLO 训练、红瓶停止、警戒区和货物异常识别。", "7.8-7.9 准备模型，7.12-7.13 部署测试。"],
        ["鸿蒙 APP 模块", "网络配置、遥控、巡检看板、环境/雷达/视觉、告警和 AI 咨询。", "贯穿全程，7.9-7.11 硬件受限时重点推进。"],
    ])

    heading(doc, "八、风险与问题处理")
    table(doc, ["风险/问题", "出现时间", "影响", "处理措施"], [
        ["小车电池无法充电，需要返工修理", "7.9", "实车调试中断，SLAM、雷达、底盘控制和视觉部署无法连续验证。", "将开发重心临时转向 APP、PPT、文档、模型训练、场景搭建和非硬件代码；硬件恢复后延长时间集中补测。"],
        ["连续暴雨导致停工和设备取用困难", "7.10-7.11", "无法稳定拿到小车，硬件联调时间进一步被压缩。", "完成中期答辩，继续整理开发材料、优化 APP 和准备最终演示方案。"],
        ["串口指向错误和多模块联调复杂", "7.13", "底盘控制、雷达和 APP 指令链路存在排查成本。", "由负责导航和控制的成员集中定位，结合实车测试逐步确认底盘控制串口和雷达串口对应关系。"],
        ["答辩前时间紧张", "7.12-7.14", "功能完善、演示稳定性和材料整理同时推进，容易出现任务冲突。", "冻结核心演示链路，将开发时间延长至 9:00-21:00，优先保障可展示功能和最终材料完整。"],
    ])

    heading(doc, "九、测试与验收")
    paragraph(doc, "测试工作围绕最终答辩展示链路展开，重点验证 APP 连接控制、巡检导航、雷达避障、视觉识别和告警展示是否能够形成闭环。")
    paragraph(doc, "连接与控制方面，团队通过配置小车 IP 和端口，验证 APP 在线状态、遥控、多档调速和停止指令是否可执行；建图与巡检方面，围绕仓储场景完成建图、巡检点设置和路线验证；雷达与视觉方面，通过设置障碍物、红瓶和警戒区，验证避障、停车、异常识别和告警提示；数据展示方面，检查环境数据、雷达状态、视觉结果和告警记录是否能够在 APP 中集中呈现。")

    heading(doc, "十、项目成果与总结")
    paragraph(doc, "经过 7 月 6 日至 7 月 15 日的集中开发，团队形成了“狭域巡航 aiPotRol：仓储环境智能巡检与货物异常监测小车”项目，完成仓储巡检场景、鸿蒙 APP、小车控制、SLAM 巡检、雷达避障、视觉识别、环境监测、告警展示以及答辩和归档材料。项目过程中虽然多次受到小车充电故障、黑屏返修和连续暴雨影响，但团队通过调整任务优先级、将等待时间转化为 PPT 和文档整理时间，并延长 7.12-7.14 开发时段，最终保证了项目收口和答辩展示。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
