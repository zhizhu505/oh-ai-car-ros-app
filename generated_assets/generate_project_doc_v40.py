from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "狭域巡航aiPotRol-项目说明书V4.0.docx"
LOGO = ROOT / "entry" / "src" / "main" / "resources" / "base" / "media" / "patrol_logo.jpg"
PLACEHOLDER_DIR = ROOT / "generated_assets" / "project_doc_placeholders"


PLACEHOLDERS = [
    ("01_cover_logo_source.png", "封面 logo 来源示意", "来自代码资源：entry/src/main/resources/base/media/patrol_logo.jpg"),
    ("02_app_home.png", "APP 首页截图占位", "对应页面：pages/Index，展示项目名、在线状态、巡检进度、重点告警与模块入口"),
    ("03_network_settings.png", "网络连接页截图占位", "对应页面：pages/NetworkSettings，展示 IP、TCP 端口、视频端口与连接状态"),
    ("04_remote_control.png", "手动遥控页截图占位", "对应页面：pages/RemoteControl，展示方向控制、摇杆、视频画面与急停"),
    ("05_mecanum_wheel.png", "麦克纳姆轮控制页截图占位", "对应页面：pages/MecanumWheel，展示四轮独立控制与横移能力"),
    ("06_patrol_dashboard.png", "巡检监控页截图占位", "对应页面：pages/PatrolDashboard，展示路线、点位、进度、雷达和视觉状态"),
    ("07_patrol_control.png", "巡检控制页截图占位", "对应页面：pages/PatrolControl，展示手动/自动模式和巡检控制按钮"),
    ("08_environment_monitor.png", "环境监测页截图占位", "对应页面：pages/EnvironmentMonitor、Module1Environment，展示温湿度、可燃气体、PM2.5、光照等数据"),
    ("09_lidar_channel.png", "雷达通道安全页截图占位", "对应页面：pages/Module3Lidar，展示前方距离、左右距离、通道占用和避障状态"),
    ("10_vision_detection.png", "视觉识别页截图占位", "对应页面：pages/Module4Vision，展示人员闯入、物体占道、红色瓶体演示和蜂鸣器控制"),
    ("11_alert_records.png", "告警记录页截图占位", "对应页面：pages/AlertRecords，展示环境、通道、视觉和巡检异常，以及 AI 分析"),
    ("12_demo_scene.png", "实物演示场景照片占位", "建议替换为最终小车、货架/管廊通道、障碍物和人员闯入演示照片"),
]


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    set_run_font(r, 16 if level == 1 else 13 if level == 2 else 11, True, "000000")
    return p


def add_para(doc, text, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_run_font(r, 10.5)
    return p


def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(f"{i}. {item}")
        set_run_font(r, 10.5)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 9, False, "666666")


def add_image(doc, path, width=5.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))


def add_small_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for idx, (left, right) in enumerate(rows):
        for cell, text, bold in [(table.cell(idx, 0), left, True), (table.cell(idx, 1), right, False)]:
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, 10, bold)
        tc_pr = table.cell(idx, 0)._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9EAF7")
        tc_pr.append(shd)
    doc.add_paragraph()


def make_placeholders():
    PLACEHOLDER_DIR.mkdir(parents=True, exist_ok=True)
    try:
        font_title = ImageFont.truetype("C:/Windows/Fonts/simhei.ttf", 34)
        font_body = ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", 22)
        font_small = ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", 18)
    except Exception:
        font_title = font_body = font_small = None
    for filename, title, note in PLACEHOLDERS:
        path = PLACEHOLDER_DIR / filename
        img = Image.new("RGB", (1200, 760), "white")
        draw = ImageDraw.Draw(img)
        draw.rectangle((30, 30, 1170, 730), outline=(48, 84, 120), width=4)
        draw.rectangle((30, 30, 1170, 120), fill=(217, 234, 247), outline=(48, 84, 120), width=4)
        draw.text((60, 55), title, fill=(0, 0, 0), font=font_title)
        draw.text((60, 170), f"图片文件名：{filename}", fill=(0, 0, 0), font=font_body)
        draw.text((60, 225), note, fill=(60, 60, 60), font=font_body)
        draw.text((60, 640), "后续替换说明：保持本文件名不变，用最终截图或实物照片覆盖此占位图。", fill=(130, 40, 40), font=font_small)
        path.write_bytes(b"")
        img.save(path)


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.add_run().add_picture(str(LOGO), width=Inches(3.0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("狭域巡航 aiPotRol")
    set_run_font(r, 24, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("仓储环境智能巡检与货物异常监测小车")
    set_run_font(r, 18, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目说明书")
    set_run_font(r, 20, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("版本：V4.0    日期：2026 年 7 月 13 日")
    set_run_font(r, 11)
    doc.add_page_break()


def build():
    make_placeholders()
    doc = Document()
    setup(doc)
    cover(doc)

    add_heading(doc, "一、项目概述", 1)
    add_para(
        doc,
        "“狭域巡航 aiPotRol”是一套面向仓储狭窄通道、货架通道和管廊式区域的智能巡检小车项目。项目以润和 iCar 智能小车为移动平台，结合 HarmonyOS 上位机 APP、TCP 网络通信、视频回传、麦克纳姆轮控制、环境监测、雷达通道安全判断、视觉异常识别和告警记录展示，形成一套能够现场演示、能够继续扩展的巡检系统。"
    )
    add_para(
        doc,
        "项目并不把目标写成泛化的大型仓储管理平台，而是聚焦小车最能发挥作用的“狭域”场景。仓储货架之间、管廊式仓储通道、设备段和高风险区域通常通道较窄，人员巡查不方便，临时堆放物、人员进入、物品倾倒、可燃气体异常、光照不足等问题容易被忽略。小车可以沿固定点位移动，在关键位置采集状态，并把异常情况直接反馈到移动端。"
    )
    add_para(
        doc,
        "当前项目已经完成了较完整的 APP 端实现。代码中包含首页、网络设置、手动遥控、麦克纳姆轮控制、巡检监控、巡检控制、环境监测、雷达通道安全、视觉识别、APP 监控和告警记录等页面。通信层支持小车 IP、TCP 指令端口和视频端口配置，巡检控制层支持启动、暂停、继续、停止、返回起点、手动/自动模式切换等操作，状态层能够统一维护车辆、环境、雷达、视觉、点位和告警记录。"
    )
    add_small_table(
        doc,
        [
            ("项目名称", "狭域巡航 aiPotRol 仓储环境智能巡检与货物异常监测小车"),
            ("软件形态", "HarmonyOS / ArkTS / ArkUI 上位机 APP + 小车端通信与巡检控制接口"),
            ("核心场景", "狭窄通道巡检、货架/管廊式区域巡检、通道占用检测、人员闯入提示、环境异常监测"),
            ("代码 logo 来源", "entry/src/main/resources/base/media/patrol_logo.jpg"),
        ],
    )

    add_heading(doc, "二、项目实现内容", 1)
    add_para(
        doc,
        "项目实现内容主要集中在上位机 APP、通信控制、巡检状态管理、环境与告警展示、雷达和视觉模块展示几个方面。APP 首页以“狭域巡航”为入口，显示在线或离线状态、当前点位、下一点位、巡检进度和重点告警；用户可以从首页进入环境感知、SLAM 导航、激光雷达、视觉识别和手动接管等模块。"
    )
    add_para(
        doc,
        "网络连接部分支持配置小车 IP、TCP 指令端口和视频端口。连接成功后，APP 会初始化消息处理器，接收小车端返回的 JSON 数据，并根据消息类型更新环境数据、巡检状态、雷达状态、视觉状态和告警记录。即使没有连接真实小车，APP 也保留离线预览能力，方便展示页面和讲解业务流程。"
    )
    add_para(
        doc,
        "手动控制部分已经实现方向控制、摇杆控制、麦克纳姆轮控制和急停等入口。底盘控制协议中包含自由控制、按键控制、四轮独立速度控制、保存图片、开始录制、结束录制、启动巡航和中止巡航等命令。麦克纳姆轮控制用于体现小车在狭窄通道中的横移和原地调整能力，这是本项目区别于普通小车遥控演示的重点之一。"
    )
    add_para(
        doc,
        "自动巡检部分已经在 APP 中实现手动模式和自动巡检模式切换，并支持开始巡检、暂停巡检、继续巡检、返回起点和停止巡检。巡检路线使用 warehouse_A 和 warehouse_B 两类路线标识，点位采用 P1 到 P6 的形式记录，包括管廊入口、A 区管线左侧、A 区管线右侧、B 区设备段、高风险设备区和返回起点。APP 可以显示当前点、下一点、巡检状态和进度，异常发生后也可以将对应点位标记为异常。"
    )
    add_para(
        doc,
        "环境监测模块围绕仓储安全常见指标设计，包含温度、湿度、可燃气体、PM2.5、光照、烟雾和气压等数据。代码中已经定义了环境快照、环境历史记录和环境阈值，页面可以显示实时数值和最近历史。环境异常按正常、预警、严重异常分级，便于在展示时说明不同风险等级。"
    )
    add_para(
        doc,
        "雷达模块用于表达通道安全能力，重点数据包括前方距离、左侧距离、右侧距离、通道状态和更新时间。通道状态分为通道畅通、减速通行、通道堵塞和通道占用。视觉模块用于表达人员闯入、物体占道和货物异常等情况，记录字段包含检测点位、是否发现人员、是否有箱体或物体进入通道、检测说明、图片地址和更新时间。"
    )
    add_para(
        doc,
        "告警模块已经实现统一记录结构，能够保存告警时间、点位、类别、消息、等级和图片地址，并预留 AI 分析结果。告警类别包括环境异常、通道占用、视觉异常和巡检异常。告警记录页支持查看单条告警分析，也支持围绕告警数据进行问答式说明，这部分体现了项目在异常处理和答辩展示上的扩展能力。"
    )

    add_heading(doc, "三、图片与演示素材占位", 1)
    add_para(
        doc,
        "本说明书已经为后续替换最新截图和实物照片预留图片文件。建议替换时保持文件名不变，直接用最终截图覆盖 generated_assets/project_doc_placeholders 目录中的同名图片，然后重新生成文档即可。下面列出的图片位对应项目中已经实现或需要展示的功能。"
    )
    for idx, (filename, title, note) in enumerate(PLACEHOLDERS[1:], 1):
        add_image(doc, PLACEHOLDER_DIR / filename, 4.8)
        add_caption(doc, f"图 3-{idx} {title}：{filename}")

    add_heading(doc, "四、功能性说明", 1)
    add_heading(doc, "4.1 上位机 APP 功能", 2)
    add_para(
        doc,
        "APP 是本项目最主要的展示和操作入口。首页需要承担总览作用，用户打开应用后应能直接看到项目名称、设备在线状态、当前巡检点位、下一巡检点位、巡检进度和近期重点告警。首页还需要提供各模块入口，使用户可以进入环境感知、SLAM 导航、激光雷达、视觉识别和手动接管页面。当前代码中的 Index 页面已经实现这些内容，并且连接弹窗支持输入小车 IP、TCP 端口和视频端口。"
    )
    add_para(
        doc,
        "网络设置功能用于解决真实小车联调问题。用户可以配置默认 IP、指令端口和视频端口，连接成功后进入在线同步状态，连接失败时给出提示。APP 还保留离线预览模式，用于在没有真实小车或网络环境不稳定时展示页面结构和业务流程。这个功能对答辩演示很重要，因为它避免了网络异常导致整个演示无法继续。"
    )
    add_para(
        doc,
        "巡检监控功能用于展示小车自动巡检的实时状态。页面应显示路线名称、当前点位、下一点位、巡检状态、巡检进度、雷达状态、视觉状态和最新告警。当前 PatrolDashboard 页面已经按这个思路实现，并使用点位卡片展示路线中的已巡检点和异常点。"
    )
    add_para(
        doc,
        "巡检控制功能用于下发自动巡检命令。用户可以在手动遥控和自动巡检之间切换，可以启动巡检、暂停巡检、继续巡检、停止巡检和返回起点。当前 PatrolControl、Module2Navigation 和 Module5AppMonitor 页面都接入了巡检控制工具，底层通过 JSON 命令向小车端发送 start_patrol、pause_patrol、resume_patrol、return_home 和 stop_patrol。"
    )

    add_heading(doc, "4.2 小车运动与手动接管功能", 2)
    add_para(
        doc,
        "手动接管是项目安全性的关键。小车在自动巡检过程中如果遇到通道堵塞、人员靠近、识别异常或导航不稳定，用户应能够立即切换到手动模式。当前项目已经实现方向按键、摇杆控制、麦克纳姆轮控制和急停入口，能够覆盖前进、后退、左平移、右平移、左旋转、右旋转、刹车停止等基本动作。"
    )
    add_para(
        doc,
        "麦克纳姆轮控制是狭域巡航场景的重要支撑。普通两轮或四轮小车在货架通道内调头和避让不方便，而麦克纳姆轮可以横向移动，也可以原地调整姿态，更适合狭窄通道、管廊式区域和货架间通道。项目中的 MecanumWheel 页面和四轮独立速度控制协议为这一能力提供了展示入口。"
    )
    add_para(
        doc,
        "视频回传和手动控制需要配合使用。用户在 RemoteControl 或 PatrolControl 页面观察小车视角后，可以根据画面判断是否继续巡检、暂停巡检或手动绕开障碍。视频端口默认独立配置，便于和指令通信分离。"
    )

    add_heading(doc, "4.3 巡检路线、点位与任务状态功能", 2)
    add_para(
        doc,
        "巡检路线采用点位化方式组织，不追求复杂的仓库全域调度，而是围绕可演示、可定位、可复盘的 P 点路线展开。当前默认点位包括 P1 管廊入口、P2 A 区管线左侧、P3 A 区管线右侧、P4 B 区设备段、P5 高风险设备区和 P6 返回起点。每个点位都包含是否已巡检和是否异常两个状态。"
    )
    add_para(
        doc,
        "任务状态包括空闲、巡检中、已暂停、正在导航、避障中、异常告警中和返回起点。APP 接收到小车端 patrol 类型消息后，会更新当前点、下一点、进度、路线和状态说明。雷达检测到通道堵塞或占用时，车辆状态可切换为避障中；告警消息指示 alerting 时，车辆状态可切换为异常告警中。"
    )

    add_heading(doc, "4.4 环境监测功能", 2)
    add_para(
        doc,
        "环境监测面向仓储安全中的常见隐患。系统需要展示温度、湿度、可燃气体、PM2.5、光照、烟雾和气压等指标，其中温度、湿度、可燃气体、PM2.5 和光照已经在模型中定义了阈值。环境数据不仅要显示当前值，还要形成历史记录，便于说明异常是在巡检过程中发现的，而不是孤立的一次展示。"
    )
    add_para(
        doc,
        "环境异常采用正常、预警、严重异常三个等级。轻微超限时 APP 可以作为预警显示，严重超限时应进入重点告警。环境记录应绑定点位，例如“P5 高风险设备区可燃气体异常”或“P3 A 区管线右侧光照不足”，这样在后续复盘时能够直接定位到异常区域。"
    )

    add_heading(doc, "4.5 雷达通道安全功能", 2)
    add_para(
        doc,
        "雷达通道安全功能主要用于判断小车前方和左右两侧是否具备安全通行条件。系统需要记录前方距离、左侧距离和右侧距离，并根据距离判断通道畅通、需要减速、通道堵塞或通道被占用。对于狭窄仓储通道，这一功能可以解释小车为什么需要低速巡检，也可以配合手动接管说明安全策略。"
    )
    add_para(
        doc,
        "当前代码已经定义 LidarSnapshot 和 ChannelStatus，并在消息处理器中支持 lidar 类型消息。收到通道堵塞或通道占用状态后，系统会同步更新雷达状态，并把车辆状态切换为避障中。后续接入真实雷达数据后，可以直接沿用现有页面和状态结构。"
    )

    add_heading(doc, "4.6 视觉识别与演示功能", 2)
    add_para(
        doc,
        "视觉识别功能用于展示人员闯入、物体占道、货物异常和特定目标检测。当前模型结构已经包含人员检测、箱体或物体进入通道、图像地址和检测说明。Module4Vision 页面还接入了红色瓶体演示相关命令，包括启动演示、停止演示、急停演示、停止蜂鸣器和获取演示状态。"
    )
    add_para(
        doc,
        "红色瓶体演示适合作为答辩中的视觉识别样例。系统可以通过 6600 端口请求演示状态，读取距离阈值、速度、是否检测到目标、蜂鸣器是否激活和更新时间。这个功能说明项目并不只是做了静态页面，而是已经为视觉识别和小车动作联动预留了实际控制链路。"
    )

    add_heading(doc, "4.7 告警记录与 AI 辅助分析功能", 2)
    add_para(
        doc,
        "告警记录用于把环境、雷达、视觉和巡检过程中的异常统一保存。每条告警包含编号、时间、点位、类别、消息、等级、图片地址和可选 AI 分析结果。首页只展示重点告警，详细记录在 AlertRecords 页面中查看。告警发生后，系统还会尝试触发 AI 分析，生成摘要、风险和处理建议。"
    )
    add_para(
        doc,
        "AI 辅助分析不是项目运行的硬性依赖，但它能增强展示效果。即使分析失败，系统也会给出网络异常提示，不影响告警记录本身。这样的设计保证了基础巡检功能稳定，同时为后续智能化处理留下扩展空间。"
    )

    add_heading(doc, "五、非功能性说明", 1)
    add_para(
        doc,
        "在易用性方面，系统应尽量让仓库管理员通过 APP 完成主要操作。连接小车、查看状态、启动巡检、暂停巡检、手动接管、查看告警都应在移动端完成，不要求普通用户进入 Linux 终端或 ROS 调试界面。页面文字应接近现场业务表达，例如当前点位、下一点位、通道占用、人员闯入、可燃气体异常等，避免只显示底层技术词。"
    )
    add_para(
        doc,
        "在实时性方面，手动控制指令需要尽量短延迟，尤其是急停、刹车和停止巡检。视频画面应能够辅助用户判断小车前方环境，环境数据、雷达状态和视觉状态应在收到消息后及时刷新。告警记录应优先展示严重异常，避免重要信息被普通日志淹没。"
    )
    add_para(
        doc,
        "在可靠性方面，系统要允许真实连接失败时继续进行页面展示。APP 已经支持离线预览，这是课堂演示和项目答辩中的必要保障。通信层解析 JSON 消息时应容忍分包和残留数据，当前 PatrolMessageHandler 已通过缓冲区方式处理按换行分割的消息。对于未知消息，系统不会直接中断运行。"
    )
    add_para(
        doc,
        "在安全性方面，手动接管、急停和停止巡检的优先级应高于普通巡检动作。小车处于狭窄通道时应优先保证低速和可控，通道堵塞、人员闯入和严重环境异常都应能够触发告警。对于红色瓶体演示和蜂鸣器控制，系统提供停止演示、急停演示和停止蜂鸣器入口，避免演示状态失控。"
    )
    add_para(
        doc,
        "在可维护性方面，项目代码已经按页面、组件、通信、状态、模型和工具类分层。页面位于 pages 目录，视频和通用 UI 位于 components 目录，TCP 连接和巡检命令位于 tcp 目录，状态集中在 PatrolStore，数据结构集中在 PatrolModels。这样的结构便于后续继续接入真实传感器、替换视觉模型、调整页面和扩展任务报告。"
    )
    add_para(
        doc,
        "在可扩展性方面，当前系统已经保留了 SLAM 导航、环境阈值、雷达通道状态、视觉检测日志、AI 告警分析、图片地址和任务状态等接口。后续可以在不大改 APP 页面结构的前提下，接入真实 ROS 节点、YOLO 推理结果、任务报告导出和历史数据查询。"
    )

    add_heading(doc, "六、概要设计", 1)
    add_para(
        doc,
        "系统整体采用上位机 APP、通信控制层、状态管理层、感知与算法层、小车硬件层的分层设计。上位机 APP 面向用户，负责页面展示和操作入口；通信控制层负责 TCP 连接、命令发送和消息解析；状态管理层负责保存车辆、环境、雷达、视觉、告警和点位数据；感知与算法层负责雷达、视觉、环境阈值和后续 SLAM 导航；硬件层负责底盘运动、传感器采集、视频回传和蜂鸣器等外设控制。"
    )
    add_para(
        doc,
        "APP 端页面结构清晰：首页负责总览，NetworkSettings 负责连接配置，RemoteControl 和 MecanumWheel 负责手动控制，PatrolDashboard 和 PatrolControl 负责巡检监控与任务控制，EnvironmentMonitor 和 Module1Environment 负责环境数据，Module3Lidar 负责通道安全，Module4Vision 负责视觉识别和红色瓶体演示，AlertRecords 负责告警记录和 AI 分析。"
    )
    add_para(
        doc,
        "通信协议分为两类。一类是底盘控制协议，使用帧格式发送自由控制、按键控制、四轮速度、保存图片、视频录制和巡航控制等命令；另一类是巡检业务 JSON 协议，用于发送 start_patrol、pause_patrol、resume_patrol、return_home、stop_patrol、set_mode 等任务指令，并接收 env、patrol、alert、lidar、vision 和 status 等类型的状态消息。"
    )
    add_para(
        doc,
        "状态管理采用集中式 Store。PatrolStore 保存当前环境快照、环境历史、小车状态、雷达状态、视觉状态、视觉日志、巡检点位和告警记录。页面通过订阅 Store 刷新视图，通信层收到消息后更新 Store。这样可以减少页面之间的数据重复，也便于在未来加入本地持久化或报告导出功能。"
    )
    add_para(
        doc,
        "巡检过程可以概括为：用户连接小车，选择手动或自动模式；自动模式下启动巡检路线，小车按点位推进；环境、雷达、视觉模块持续或按点位返回状态；出现异常时写入告警记录并更新点位状态；用户根据视频和告警决定继续巡检、暂停巡检、返回起点或手动接管。"
    )

    add_heading(doc, "七、后续补充与完善方向", 1)
    add_para(
        doc,
        "后续文档中需要优先补充最终 APP 截图和实物演示图片。本文档已经预留了对应图片文件名，包括首页、网络设置、手动遥控、麦克纳姆轮、巡检监控、巡检控制、环境监测、雷达通道安全、视觉识别、告警记录和实物演示场景。替换图片后，文档的展示完整度会明显提升。"
    )
    add_para(
        doc,
        "功能上可以继续完善真实传感器接入、SLAM 地图和点位配置、YOLO 模型识别、任务报告导出、历史记录查询和断线重连。当前 APP 层已经具备较好的结构，后续工作重点应放在真实数据质量和演示稳定性上。"
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("狭域巡航 aiPotRol 项目说明书 V4.0")
    set_run_font(r, 9, False, "888888")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
