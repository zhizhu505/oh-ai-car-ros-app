from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SRC = r"E:\test_java\huawei\项目开发及管理文档_狭域巡航_aiPotRol_修改前备份.docx"
OUT = r"E:\test_java\huawei\项目开发及管理文档_狭域巡航_aiPotRol_竖屏更新版.docx"
FONT = "宋体"


def set_text(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    apply_run_style(r)


def apply_run_style(run, size=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size:
        run.font.size = Pt(size)


def style_paragraphs(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            apply_run_style(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply_run_style(run)


def compact_table(table, size):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    apply_run_style(run, size=size)


def build():
    doc = Document(SRC)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.6)

    # 每日开发记录：补入 7.14 新增实际情况。
    daily = doc.tables[4]
    row_714 = daily.rows[9].cells
    set_text(row_714[2], "上午小车出现黑屏并再次无法充电，团队将小车送去维修；等待维修期间继续整理结题 PPT、完善文档和最终演示流程，并对已完成的 APP、视觉、雷达和导航内容做答辩材料梳理。")
    set_text(row_714[3], "根据硬件状态再次调整节奏，把无法实车调试的时间转为答辩材料收口和演示逻辑梳理，尽量降低硬件不稳定对最终答辩的影响。")

    # 进度偏差与应对：在原有三段基础上补充 7.14 黑屏送修信息。
    doc.paragraphs[26].text = (
        "原计划中，7 月 7 日至 7 月 11 日应是小车硬件持续联调的主要阶段，但 7 月 9 日小车电池无法充电导致硬件开发被迫中断，"
        "7 月 10 日和 7 月 11 日又遇到连续暴雨，进一步影响设备取还和实车调试。7 月 14 日上午小车又出现黑屏并再次无法充电，"
        "团队将小车送去维修，答辩前硬件验证时间再次被压缩。"
    )
    doc.paragraphs[28].text = (
        "从 7 月 11 日到 7 月 14 日，团队将开发时间延长为每天 9:00-21:00，利用连续长时段补齐实车调试不足。"
        "其中 7 月 14 日维修等待时间主要用于结题 PPT 整理、文档收口和演示流程梳理，保证最终答辩材料与项目过程记录保持完整。"
    )

    # 风险表：增加 7.14 新风险，保留原版表格风格。
    risk = doc.tables[7]
    new = risk.add_row().cells
    values = [
        "小车黑屏并再次无法充电",
        "7.14 上午",
        "答辩前实车验证和最终演示排练再次受到影响。",
        "立即送修小车；等待期间集中整理结题 PPT、完善文档和演示流程，优先保证答辩材料完整。",
    ]
    for cell, value in zip(new, values):
        set_text(cell, value)

    # 总结：补入新事件，但保留原总结口吻。
    doc.paragraphs[41].text = (
        "从管理角度看，本项目最大的特点是实际开发过程受到硬件故障和天气影响，但团队通过及时调整任务优先级，将硬件不可用时间转化为软件、"
        "文档、模型和展示材料的推进时间。7 月 14 日上午小车黑屏并再次无法充电后，团队将小车送去维修，同时继续整理结题 PPT 和演示逻辑。"
        "这些调整保证了短周期实训中计划、执行、偏差、调整和收口的全过程能够被完整记录。"
    )

    style_paragraphs(doc)
    # 竖屏下宽表需要略微压缩字号，避免版面过宽。
    compact_table(doc.tables[3], 7.5)
    compact_table(doc.tables[4], 8.5)
    compact_table(doc.tables[7], 8.8)
    compact_table(doc.tables[8], 8.8)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
