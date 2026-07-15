from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"E:\test_java\huawei\项目开发及管理文档_面向地下综合管廊智能巡检小车.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_paragraph(doc, text="", first_line=True, size=10.5, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(16 if level == 1 else 13)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
        if widths:
            table.rows[0].cells[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), align=WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def add_gantt(doc):
    dates = ["7.6", "7.7", "7.8", "7.9", "7.10", "7.11", "7.12", "7.13", "7.14", "7.15"]
    tasks = [
        ("企业授课与任务理解", [0]),
        ("开发环境与小车连接", [0, 1]),
        ("硬件熟悉与基础控制", [1, 2]),
        ("传感器与环境感知开发", [2, 3, 6, 7]),
        ("SLAM 建图与定点巡检", [2, 3, 6, 7, 8]),
        ("雷达追踪、避障与串口排查", [2, 3, 6, 7, 8]),
        ("视觉数据集、YOLO 训练与识别", [2, 3, 6, 7]),
        ("鸿蒙 APP 页面、通信与告警", [3, 4, 5, 6, 7, 8]),
        ("文档、PPT 与答辩材料", [3, 4, 5, 6, 7, 8, 9]),
        ("中期答辩", [4]),
        ("实车集中联调与结题演示", [6, 7, 8]),
        ("最终答辩", [9]),
    ]
    table = doc.add_table(rows=1, cols=1 + len(dates))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(table.rows[0].cells[0], "任务", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.rows[0].cells[0], "D9EAF7")
    for i, d in enumerate(dates, start=1):
        set_cell_text(table.rows[0].cells[i], d, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for task, active in tasks:
        cells = table.add_row().cells
        set_cell_text(cells[0], task, size=9)
        for i in range(len(dates)):
            text = "■" if i in active else ""
            set_cell_text(cells[i + 1], text, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            if i in active:
                set_cell_shading(cells[i + 1], "70AD47")
            if i in [3, 4, 5]:
                if i not in active:
                    set_cell_shading(cells[i + 1], "FCE4D6")
    add_paragraph(doc, "说明：绿色表示实际投入开发或交付的主要时段，浅橙色表示受小车电池故障、连续暴雨和设备不可用影响，需要调整计划的时段。7.12 至 7.14 为集中补强阶段，团队将开发时间延长至 9:00-21:00，用于补齐实车联调和最终展示能力。")


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.2)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("项目开发及管理文档")
    r.bold = True
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("面向地下综合管廊的智能巡检与异常监测小车")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(15)

    add_heading(doc, "一、项目基本信息")
    add_table(
        doc,
        ["项目项", "内容"],
        [
            ["项目名称", "面向地下综合管廊的智能巡检与异常监测小车"],
            ["项目周期", "2026 年 7 月 6 日至 2026 年 7 月 15 日"],
            ["最终答辩", "2026 年 7 月 15 日全天线下结题答辩"],
            ["团队成员", "武芷竹、胡雅欣、周诗晴、曹玥"],
            ["应用场景", "地下综合管廊、狭窄检修通道、设备管线巡检场景"],
            ["核心硬件", "智能小车、激光雷达、深度/视觉相机、环境感知传感器"],
            ["软件环境", "Ubuntu、ROS、YOLO 目标检测、HarmonyOS / ArkTS / ArkUI、TCP 通信"],
            ["主要交付", "鸿蒙 APP、巡检演示场景、实车巡检与避障能力、视觉识别能力、项目文档、测试报告、使用手册、答辩 PPT 和演示视频"],
        ],
        widths=[Cm(3.5), Cm(12)],
    )

    add_heading(doc, "二、项目背景与开发目标")
    add_paragraph(doc, "2026 年 7 月 6 日，企业老师围绕智能小车、ROS、激光雷达、视觉识别、建图导航和上位机控制等内容进行授课与答疑。团队在当天明确本次实训不是单一功能实验，而是需要围绕智能硬件形成完整的场景化项目，并在 7 月 15 日最终答辩中展示项目立项、开发过程、成员分工、设计实现和实际成果。")
    add_paragraph(doc, "结合课程要求和现有硬件条件，团队将项目场景确定为地下综合管廊巡检。该场景具有空间狭窄、光照不足、空气质量风险、人工巡检成本高和异常点位需要及时定位等特点，能够自然承接小车导航、雷达避障、环境感知、视觉识别和鸿蒙 APP 监控等模块。")
    add_paragraph(doc, "项目目标是以智能小车为移动平台，通过 ROS 完成底盘控制、SLAM 建图和巡检导航，通过传感器采集管廊环境数据，通过激光雷达完成障碍检测和避障，通过视觉识别完成人员闯入、红瓶/异物等异常目标检测，并通过鸿蒙 APP 完成状态展示、远程控制、告警记录和 AI 咨询辅助。")

    add_heading(doc, "三、需求分析")
    add_heading(doc, "3.1 功能需求", 2)
    add_table(
        doc,
        ["序号", "需求项", "实际落地方式"],
        [
            ["1", "管廊环境数据采集与异常告警", "围绕温湿度、PM2.5、光照、GPS 等传感器完成代码构建和数据组织，并在 APP 中通过数据卡片和图表展示异常状态。"],
            ["2", "SLAM 建图与固定点巡检", "基于实际搭建场景生成地图，设计巡检路线和定点巡检点位，调整成本地图参数以支持狭窄通道通行。"],
            ["3", "激光雷达避障与通道堵塞检测", "完成雷达追踪、自动避障、通道占用检测和相关串口排查，使小车能够在管廊通道中识别障碍并采取减速或停止策略。"],
            ["4", "视觉识别人员闯入、红瓶/异物等异常目标", "构建目标检测数据集，训练 YOLO 模型，完成红瓶停止、警戒区识别和场景化异常目标识别。"],
            ["5", "鸿蒙 APP 状态展示、数据展示、告警和控制", "实现网络配置、遥控控制、巡检看板、环境监测、雷达状态、视觉识别、告警记录、AI 咨询和横竖屏适配等页面与功能。"],
        ],
    )
    add_heading(doc, "3.2 非功能需求", 2)
    add_paragraph(doc, "项目需要满足演示稳定性、页面可读性、现场操作可控性和文档可追溯性要求。由于最终答辩采用现场演示与成果汇报结合的形式，系统不仅要能运行核心功能，还需要在答辩场景中清楚体现开发过程、问题处理、成员贡献和最终效果。")
    add_paragraph(doc, "在实际开发中，团队还需要考虑硬件可用时间受限、天气导致取还设备困难、小车电池和串口等问题带来的进度风险。因此项目管理上采用“硬件可用时优先实车联调，硬件不可用时推进 APP、模型、文档、PPT 和场景搭建”的策略，保证总进度不断档。")

    add_heading(doc, "四、总体设计")
    add_paragraph(doc, "系统整体由智能小车硬件层、ROS 控制与导航层、传感器与视觉感知层、TCP 通信层、鸿蒙 APP 展示控制层和答辩演示场景组成。智能小车负责运动执行，ROS 负责底盘控制、建图导航和节点调度，传感器与雷达提供环境和障碍信息，视觉识别模块提供异常目标判断，鸿蒙 APP 负责远程连接、状态展示、控制指令下发和告警汇总。")
    add_table(
        doc,
        ["层级", "组成", "作用"],
        [
            ["硬件执行层", "智能小车、底盘、电池、雷达、相机、传感器", "完成移动、采集、识别和现场演示所需的基础硬件能力。"],
            ["ROS 与算法层", "SLAM、导航、雷达处理、视觉识别、传感器数据处理", "完成建图、巡检路线执行、避障、异常判断和数据输出。"],
            ["通信层", "TCP 控制通道、数据消息处理", "连接鸿蒙 APP 与小车端，实现控制命令发送和状态信息接收。"],
            ["APP 展示层", "网络配置、巡检看板、遥控、环境、雷达、视觉、告警和 AI 咨询页面", "为现场操作和答辩展示提供统一入口。"],
            ["场景与交付层", "管廊布局、PPT、文档、测试报告、使用手册、演示视频", "支撑最终答辩和课程归档。"],
        ],
    )

    add_heading(doc, "五、项目管理主线与时间推进")
    add_paragraph(doc, "本项目的核心管理难点在于开发周期短、硬件依赖强、天气和设备状态不可控。团队实际从 7 月 6 日开始正式实践，到 7 月 15 日全天答辩，完整开发窗口不足十天。因此项目管理重点不是简单列出模块，而是每天根据设备状态和答辩节点动态调整工作重心。")
    add_heading(doc, "5.1 项目进度甘特图", 2)
    add_gantt(doc)

    add_heading(doc, "5.2 每日开发与管理记录", 2)
    add_table(
        doc,
        ["日期", "项目状态", "当天重点工作", "管理调整与结果"],
        [
            ["7.6 周一", "正式启动", "企业老师授课，团队明确实训内容、课程要求和最终答辩目标；完成 Ubuntu、ROS 相关虚拟机环境搭建；初步确定以智能小车为核心进行场景化开发。", "建立团队协作方式，明确需要从单点实验转向完整项目；将后续任务拆分为环境感知、SLAM 导航、雷达避障、视觉识别和鸿蒙 APP。"],
            ["7.7 周二", "硬件上手", "下午拿到小车后开始熟悉硬件功能，完成小车实机环境配置、电脑连接小车、APP 命令运行小车控制系统和基础控制验证。", "优先处理硬件连通问题，确认小车、虚拟机和 APP 控制链路具备继续开发的基础条件。"],
            ["7.8 周三", "多模块并行", "推进温湿度传感器代码、实车 SLAM 建图、目标检测数据集构建、雷达追踪与自动避障配置、小程序/APP 内容设计。", "进入并行开发阶段，各成员围绕自己的模块产出可验证结果，同时为中期检查准备材料。"],
            ["7.9 周四", "硬件受阻", "小车电池出现充不上电问题，需要返工修理，实车调试受限；团队转向 GPS、PM2.5、光照传感器代码构建、YOLO 训练、SLAM 自动导航方案、鸿蒙 APP 调试、需求文档和中期 PPT。", "及时调整计划，减少对硬件的等待，把不可依赖小车的工作前置完成，包括文档、PPT、场景设计、模型训练和 APP 源码调试。"],
            ["7.10 周五", "中期节点", "连续暴雨影响设备取用，完成中期答辩；继续修改中期 PPT，整理前期工作、后续规划和鸿蒙适配小车上位机 APP 代码。", "以中期答辩倒逼阶段总结，重新确认剩余开发重点：实车联调、APP 展示、雷达/视觉/导航闭环。"],
            ["7.11 周六", "持续受限", "暴雨影响仍然存在，无法稳定拿到小车；团队继续补充 APP、文档、演示方案、场景搭建和后续联调计划。", "将受阻时间转化为材料和软件完善时间，为恢复硬件后快速联调做准备。"],
            ["7.12 周日", "集中补强", "延长开发时间至 9:00-21:00，推进小车运动连接控制、视频修复、界面功能整合、成本地图参数调整、定点巡检测试、AI 大模型接入、横竖屏适配、视觉环境配置和场景代码。", "进入冲刺阶段，管理重点转为把分散模块串成可演示流程，优先保障最终答辩能稳定展示。"],
            ["7.13 周一", "联调深化", "继续 9:00-21:00 集中开发，优化 APP 页面布局以适配地下管廊场景，配置传感器数据和数据面板图表，实现机器人转向控制及多档调速；完成红瓶停止部署、警戒区识别模型训练、AI 咨询、雷达功能接入、地图生成、巡检路线设计和串口问题排查。", "从功能可用转向展示完整度，重点解决界面、数据、控制、雷达、视觉和导航之间的联动问题。"],
            ["7.14 周二", "答辩前收口", "继续集中完善小车实车巡检、避障、视觉识别、APP 告警展示、文档和结题 PPT，对最终演示流程进行排练和风险检查。", "冻结主要演示链路，减少大改动，集中修复影响现场展示的问题。"],
            ["7.15 周三", "最终答辩", "全天线下结题答辩，围绕项目立项、开发过程、成员贡献、设计实现、测试结果和最终成果进行展示。", "以项目过程管理和实际成果为汇报主线，突出短周期、强约束条件下的计划调整和团队协作。"],
        ],
    )

    add_heading(doc, "5.3 进度偏差与应对")
    add_paragraph(doc, "原计划中，7 月 7 日至 7 月 11 日应是小车硬件持续联调的主要阶段，但 7 月 9 日小车电池无法充电导致硬件开发被迫中断，7 月 10 日和 7 月 11 日又遇到连续暴雨，进一步影响设备取还和实车调试。该偏差直接压缩了实车验证时间。")
    add_paragraph(doc, "团队没有等待硬件恢复，而是将工作拆为“必须依赖小车”和“不依赖小车”两类。硬件不可用时，优先完成鸿蒙 APP 页面、通信逻辑、告警展示、AI 咨询、视觉模型训练、文档撰写、PPT 制作和管廊场景搭建；硬件恢复后，则集中进行小车控制、SLAM 导航、雷达避障、视觉识别和 APP 联调。")
    add_paragraph(doc, "从 7 月 11 日到 7 月 14 日，团队将开发时间延长为每天 9:00-21:00，利用连续长时段补齐实车调试不足。这一调整保证了最终答辩前能够形成可演示的系统流程，而不是只停留在单个模块或文档材料。")

    add_heading(doc, "六、成员分工与贡献")
    add_table(
        doc,
        ["成员", "角色定位", "主要贡献"],
        [
            ["武芷竹", "组长 / SLAM 导航 / 项目统筹", "负责项目进度协调、SLAM 建图与自动导航验证、管廊巡检路线设计、成本地图参数调整、定点巡检测试、串口问题排查和多模块联调推进。"],
            ["胡雅欣", "鸿蒙 APP / 环境感知 / 展示材料", "负责鸿蒙 APP 页面和小车控制适配、环境传感器代码构建、数据面板和图表开发、机器人转向与多档调速、需求文档和结题展示 PPT 整理。"],
            ["周诗晴", "雷达避障 / AI 咨询 / APP 功能接入", "负责雷达追踪与自动避障配置、雷达相关功能接入 APP、AI 咨询功能完善、横竖屏适配、logo 和界面细节完善。"],
            ["曹玥", "视觉识别 / YOLO 模型 / 场景代码", "负责目标检测数据集构建、YOLO 模型训练、红瓶停止功能部署、警戒区识别模型训练、两种场景下的小车行动代码和视觉功能环境配置。"],
        ],
    )
    add_paragraph(doc, "分工并非固定不变。中期前团队按照模块划分任务，中期后根据硬件故障、暴雨停工和最终答辩压力进行动态调整：硬件恢复时集中支持实车联调，硬件不可用时共同推进文档、PPT、APP 和场景布置，保证每个时间段都有可交付成果。")

    add_heading(doc, "七、模块设计与实现概况")
    add_table(
        doc,
        ["模块", "开发内容", "与时间进度的关系"],
        [
            ["环境感知模块", "完成温湿度、PM2.5、光照、GPS 等代码构建和数据组织，在鸿蒙 APP 中进行图表化展示。", "7.8 开始推进，7.9 硬件受阻时继续完成代码与文档，7.13 进一步完善数据面板。"],
            ["SLAM 导航模块", "完成实车 SLAM 建图、自动导航、巡检点设计、地图生成、路线规划和窄道通行参数调整。", "7.8-7.9 初步验证，7.12-7.14 作为实车集中联调重点。"],
            ["雷达避障模块", "完成雷达追踪、自动避障、通道堵塞检测、警戒区域设计和串口指向错误排查。", "7.8 开始配置，7.12-7.13 与 APP、底盘控制和现场路线进行联调。"],
            ["视觉识别模块", "完成目标检测数据集构建、YOLO 模型训练、红瓶停止、警戒区识别和人员/异物异常目标识别。", "7.8-7.9 完成数据集与模型准备，7.12-7.13 部署到实车场景。"],
            ["鸿蒙 APP 模块", "完成网络配置、首页、遥控控制、巡检看板、环境监测、雷达页面、视觉页面、告警记录、AI 咨询和横竖屏适配。", "贯穿整个项目，尤其在 7.9-7.11 硬件受限期间承担主要可推进工作，7.12 后用于联调展示。"],
        ],
    )

    add_heading(doc, "八、风险与问题处理")
    add_table(
        doc,
        ["风险/问题", "出现时间", "影响", "处理措施"],
        [
            ["小车电池无法充电，需要返工修理", "7.9", "实车调试中断，SLAM、雷达、底盘控制和视觉部署无法连续验证。", "将开发重心临时转向 APP、PPT、文档、模型训练、场景搭建和非硬件代码；硬件恢复后延长时间集中补测。"],
            ["连续暴雨导致停工和设备取用困难", "7.10-7.11", "无法稳定拿到小车，硬件联调时间进一步被压缩。", "完成中期答辩，继续整理开发材料、优化 APP 和准备最终演示方案。"],
            ["串口指向错误和多模块联调复杂", "7.13", "底盘控制、雷达和 APP 指令链路存在排查成本。", "由负责导航和控制的成员集中定位，结合实车测试逐步确认底盘控制串口和雷达串口对应关系。"],
            ["答辩前时间紧张", "7.12-7.14", "功能完善、演示稳定性和材料整理同时推进，容易出现任务冲突。", "冻结核心演示链路，将开发时间延长至 9:00-21:00，优先保障可展示功能和最终材料完整。"],
        ],
    )

    add_heading(doc, "九、测试与验收")
    add_paragraph(doc, "测试工作围绕最终答辩展示链路展开，重点验证 APP 能否连接小车、控制指令能否发送、环境和告警数据能否展示、雷达和视觉结果能否支撑异常监测演示，以及小车能否在管廊场景中完成巡检与避障。")
    add_table(
        doc,
        ["测试项", "测试内容", "验收标准"],
        [
            ["APP 与小车连接测试", "输入小车 IP、TCP 端口和视频端口，建立控制连接。", "APP 能显示在线状态，并进入首页、遥控和巡检页面。"],
            ["遥控与多档速度控制测试", "使用摇杆、方向按钮和多档调速控制小车运动。", "小车能执行前进、后退、转向、停止和速度调整。"],
            ["SLAM 建图与定点巡检测试", "在管廊场景中生成地图，设置巡检点并执行路线。", "小车能够按照路线到达巡检点，并在窄道中保持可控。"],
            ["雷达避障测试", "在通道中设置障碍物，验证雷达追踪、避障和停车策略。", "检测到障碍时能够减速、停止或提示通道占用。"],
            ["视觉红瓶/警戒区识别测试", "在场景中放置红瓶或设置警戒区域，验证 YOLO 识别结果。", "识别到目标后能够触发停止或异常提示。"],
            ["告警展示与数据面板测试", "查看环境数据、雷达状态、视觉状态和告警记录。", "APP 能集中展示状态、异常类型、时间和处理建议。"],
        ],
    )

    add_heading(doc, "十、项目成果与总结")
    add_paragraph(doc, "经过 7 月 6 日至 7 月 15 日的集中开发，团队形成了以地下综合管廊为场景的智能巡检小车项目。项目成果包括管廊巡检场景布局、鸿蒙 APP、小车巡检与远程控制能力、SLAM 建图与定点巡检能力、雷达避障能力、视觉识别能力、环境监测与告警展示能力，以及中期答辩、结题答辩、项目文档、测试报告、使用手册和演示视频等材料。")
    add_paragraph(doc, "从管理角度看，本项目最大的特点是实际开发过程受到硬件故障和天气影响，但团队通过及时调整任务优先级，将硬件不可用时间转化为软件、文档、模型和展示材料的推进时间，并在 7 月 12 日至 7 月 14 日通过延长开发时段集中完成实车联调。最终文档重点记录的不只是系统实现结果，也包括短周期实训中计划、执行、偏差、调整和收口的全过程。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
